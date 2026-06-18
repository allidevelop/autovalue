"""Імпорт аналогів із готових звітів оцінки (`.doc`) у базу `report_comparables`.

Детермінований конвеєр (без LLM): LibreOffice конвертує `.doc → .docx`, python-docx
читає таблицю-порівняння (5 аналогів). Поля нормалізуються й зберігаються через
`report_db.upsert_many` (дедуп між звітами). Зберігаємо ЛИШЕ аналоги (не сам об'єкт).
"""

from __future__ import annotations

import argparse
import json
import re
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from datetime import date, datetime
from pathlib import Path
from typing import Any

from realtify import report_db
from realtify.analog_cache import address_key
from realtify.excel_tools import libreoffice_path

_UA_MONTHS = {
    "січня": 1, "лютого": 2, "березня": 3, "квітня": 4, "травня": 5, "червня": 6,
    "липня": 7, "серпня": 8, "вересня": 9, "жовтня": 10, "листопада": 11, "грудня": 12,
}


@dataclass
class ParsedReport:
    report_id: str
    rows: list[dict[str, Any]] = field(default_factory=list)
    complexes: set[str] = field(default_factory=set)
    warnings: list[str] = field(default_factory=list)


def import_reports(inputs: list[Path], *, dry_run: bool = False, progress=None) -> dict[str, Any]:
    docs = _gather_docs(inputs)
    if not docs:
        return {"error": "no .doc files found", "inputs": [str(p) for p in inputs]}

    summary: dict[str, Any] = {
        "reports_total": len(docs), "reports_ok": 0, "reports_failed": 0,
        "rows_parsed": 0, "buildings": 0, "complexes": 0, "failed_files": [], "warnings": [],
    }
    all_rows: list[dict[str, Any]] = []
    buildings: set[str] = set()
    complexes: set[str] = set()

    with tempfile.TemporaryDirectory(prefix="report_import_") as tmp_name:
        tmp = Path(tmp_name)
        docx_map = _convert_all(docs, tmp)
        for i, doc in enumerate(docs, start=1):
            docx = docx_map.get(doc)
            if docx is None:
                summary["reports_failed"] += 1
                summary["failed_files"].append(f"{doc.name}: convert failed")
                continue
            try:
                parsed = parse_report(docx, report_id=_report_id(doc, i))
                if not parsed.rows:
                    summary["reports_failed"] += 1
                    summary["failed_files"].append(f"{doc.name}: no comparables parsed")
                    continue
                all_rows.extend(parsed.rows)
                buildings.update(r["address_key"] for r in parsed.rows if r.get("address_key"))
                complexes.update(parsed.complexes)
                summary["reports_ok"] += 1
                summary["warnings"].extend(f"{doc.name}: {w}" for w in parsed.warnings)
            except Exception as exc:  # noqa: BLE001 — один битий звіт не валить весь імпорт
                summary["reports_failed"] += 1
                summary["failed_files"].append(f"{doc.name}: {type(exc).__name__}: {exc}")
            if progress and i % 10 == 0:
                progress(f"Оброблено {i}/{len(docs)} звітів…")

    summary["rows_parsed"] = len(all_rows)
    summary["buildings"] = len(buildings)
    summary["complexes"] = len(complexes)
    if not dry_run and all_rows:
        summary["db"] = report_db.upsert_many(all_rows)
    elif dry_run:
        summary["db"] = {"dry_run": True}
    return summary


def parse_report(docx_path: Path, *, report_id: str) -> ParsedReport:
    from docx import Document

    document = Document(str(docx_path))
    full_text = "\n".join(p.text for p in document.paragraphs)
    for table in document.tables:
        full_text += "\n" + "\n".join(c.text for row in table.rows for c in row.cells)

    city = _extract_city(full_text)
    report_date = _extract_report_date(full_text)
    table = _find_comparables_table(document.tables)
    parsed = ParsedReport(report_id=report_id)
    if table is None:
        parsed.warnings.append("comparables_table_not_found")
        return parsed

    labels = _row_labels(table)
    n_cols = max((len(r.cells) for r in table.rows), default=0)
    # Колонки 1..n-2 = аналоги; остання колонка = об'єкт оцінки (пропускаємо).
    comparable_cols = list(range(1, max(1, n_cols - 1)))
    rasters = _extract_analog_images(docx_path)  # [(bytes, listing_url), ...] у порядку документа
    url_index: dict[str, int] = {}
    for i, (_data, url) in enumerate(rasters):
        lid = _listing_id(url)
        if lid:
            url_index.setdefault(lid, i)
    for col in comparable_cols:
        rec = _build_record(table, labels, col, report_id, city, report_date)
        if rec is None:
            continue
        _attach_screenshot(rec, rasters, url_index, col)
        parsed.rows.append(rec)
        if rec.get("complex_name"):
            parsed.complexes.add(str(rec["complex_name"]).lower())
    if not parsed.rows:
        parsed.warnings.append("no_valid_comparable_columns")
    return parsed


# Сигнали, що растр — це скриншот оголошення-аналога (а не вступне фото/мапа/лого).
# Перевіряється текст/гіперпосилання ОДРАЗУ ПІСЛЯ зображення (підпис-знизу).
_LISTING_SIGNALS = (
    "rieltor", "dom.ria", "//ria.com", "olx.ua", "/olx", "lun.ua", "real-estate.lviv",
    "m2bomber", "flatfy", "realty-prodaja", "продаж квартир", "оголошення №", "obyavlenie", "objava",
)


def _extract_analog_images(docx_path: Path, *, min_size: int = 40_000, limit: int = 8) -> list[tuple[bytes, str]]:
    """Скриншоти аналогів зі звіту — (байти, URL_оголошення) у порядку документа.
    Аналог визначається не розміром/позицією (перші великі растри — це вступні
    фото/мапа), а ЛІСТИНГ-СИГНАЛОМ у підписі під зображенням (домен оголошень або
    «Продаж квартир»/«Оголошення №»). URL з підпису дає змогу прив'язувати скрин до
    колонки за ID оголошення (а не лише позиційно). Корпус: 5/звіт у 184/187."""
    import zipfile

    out: list[tuple[bytes, str]] = []
    try:
        with zipfile.ZipFile(docx_path) as z:
            doc = z.read("word/document.xml").decode("utf-8", "ignore")
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8", "ignore")
            rid2t = dict(re.findall(r'Id="([^"]+)"[^>]*Target="([^"]+)"', rels))
            # Потік токенів: зображення / гіперпосилання / текст — кожному растру
            # приписуємо текст, що йде ОДРАЗУ ПІСЛЯ нього (до наступного растру).
            rasters: list[list[str]] = []  # [arc, following_text]
            token = re.compile(
                r'<a:blip[^>]*r:embed="([^"]+)"|<w:hyperlink[^>]*r:id="([^"]+)"|<w:t[^>]*>(.*?)</w:t>'
            )
            for m in token.finditer(doc):
                emb, hl, txt = m.group(1), m.group(2), m.group(3)
                if emb:
                    tgt = rid2t.get(emb, "")
                    if tgt.lower().split("/")[-1].endswith((".png", ".jpg", ".jpeg")):
                        arc = tgt if tgt.startswith("word/") else "word/" + tgt.lstrip("./")
                        rasters.append([arc, ""])
                elif hl and rasters:
                    rasters[-1][1] += " " + rid2t.get(hl, "")
                elif txt and rasters:
                    rasters[-1][1] += " " + re.sub(r"<[^>]+>", "", txt)

            for arc, following in rasters:
                if not any(sig in following.lower()[:220] for sig in _LISTING_SIGNALS):
                    continue
                try:
                    data = z.read(arc)
                except KeyError:
                    continue
                if len(data) < min_size:
                    continue
                m = re.search(r"https?://[^\s<>\"]+", following[:300])
                out.append((data, m.group(0) if m else ""))
                if len(out) >= limit:
                    break
    except Exception:  # noqa: BLE001 — відсутність скринів не валить імпорт даних
        return out
    return out


def _listing_id(url: str | None) -> str:
    """ID оголошення з URL (остання послідовність ≥5 цифр у шляху, без query)."""
    if not url:
        return ""
    path = str(url).split("?")[0]
    nums = re.findall(r"\d{5,}", path)
    return nums[-1] if nums else ""


def _attach_screenshot(
    rec: dict[str, Any], rasters: list[tuple[bytes, str]], url_index: dict[str, int], col: int
) -> None:
    """Прив'язка скрина до аналога: спершу за ID оголошення (URL колонки ↔ URL у
    підписі скрина), якщо колонка має реальний URL; інакше — позиційно (col-1).
    Це виправляє нечасту перестановку скринів між колонками (корпус: 8/652)."""
    src = str(rec.get("source_url") or "")
    idx: int | None = None
    if src.startswith("http") and "report.local" not in src:
        cid = _listing_id(src)
        idx = url_index.get(cid)  # точний матч за ID оголошення
        if idx is None:
            # Без точного матчу: позиційний фолбек дозволено ЛИШЕ якщо растр на цій
            # позиції не має ВЛАСНОГО (іншого) ID. Якщо в нього явно інший ID — це
            # чуже оголошення, скрин не ставимо (щоб не показати іншу ціну клієнту).
            pos = col - 1
            if 0 <= pos < len(rasters):
                pos_id = _listing_id(rasters[pos][1])
                if not pos_id or pos_id == cid:
                    idx = pos
    else:
        idx = col - 1  # синтетичний URL (підпис без посилання) — позиційний фолбек (~99%)
    if idx is None or not 0 <= idx < len(rasters):
        return
    data, _url = rasters[idx]
    ak = address_key(
        city=rec.get("city"), address=rec.get("address"),
        property_type=rec.get("property_type") or "apartment", complex_name=None,
    )
    dk = report_db.compute_dedup_key(ak, rec.get("area_m2"), rec.get("price_usd"), rec.get("floor_or_level"))
    try:
        rec["screenshot_path"] = str(report_db.store_screenshot(dk, data))
    except Exception:  # noqa: BLE001
        pass


# ── парсинг таблиці ──────────────────────────────────────────────────────────

def _find_comparables_table(tables: list[Any]) -> Any | None:
    for table in tables:
        text = " ".join(c.text.lower() for row in table.rows for c in row.cells)
        if "адреса об" in text and "ціна пропозиції" in text and "джерело" in text:
            return table
    # запасний варіант: за двома мітками
    for table in tables:
        text = " ".join(c.text.lower() for row in table.rows for c in row.cells)
        if "площа" in text and "ціна пропозиції" in text:
            return table
    return None


def _row_labels(table: Any) -> dict[str, int]:
    """field_name → row_index, за вмістом першої колонки."""
    out: dict[str, int] = {}
    for ri, row in enumerate(table.rows):
        if not row.cells:
            continue
        label = row.cells[0].text.strip().lower()
        if not label:
            continue
        field_name = _classify_label(label)
        if field_name and field_name not in out:
            out[field_name] = ri
    return out


def _classify_label(label: str) -> str | None:
    has_price = "ціна" in label and "пропозиц" in label
    per_m2 = "1м2" in label or "1 м2" in label or "м2" in label or "кв.м" in label or "м²" in label
    if "адреса" in label:
        return "address"
    if "площа" in label:
        return "area_m2"
    if "поверх" in label:
        return "floor_or_level"
    if has_price and per_m2:
        return "price_per_m2_usd"
    if has_price:
        return "price_usd"
    if "розташуванн" in label or "місце" in label:
        return "location_quality"
    if "клас" in label:
        return "building_class"
    if "оздоблен" in label:
        return "condition"
    if "термін" in label:
        return "delivery_date"
    if "джерело" in label:
        return "source_url"
    return None


def _build_record(
    table: Any, labels: dict[str, int], col: int, report_id: str,
    city: str | None, report_date: date | None,
) -> dict[str, Any] | None:
    def cell(field_name: str) -> str | None:
        ri = labels.get(field_name)
        if ri is None or ri >= len(table.rows):
            return None
        cells = table.rows[ri].cells
        if col >= len(cells):
            return None
        text = cells[col].text.strip()
        return text or None

    raw_addr = cell("address")
    area = _parse_number(cell("area_m2"))
    price = _parse_money_usd(cell("price_usd"))
    if not raw_addr or area is None or price is None:
        return None  # без критичних полів аналог не годиться для оцінки

    address, complex_name = _split_address_complex(raw_addr)
    ppm = _parse_money_usd(cell("price_per_m2_usd"))
    if ppm is None and price and area:
        ppm = round(price / area, 2)
    src = _clean_url(cell("source_url"))
    return {
        "city": city,
        "address": address,
        "complex_name": complex_name,
        "property_type": "apartment",
        "area_m2": area,
        "price_usd": price,
        "price_per_m2_usd": ppm,
        "floor_or_level": cell("floor_or_level"),
        "rooms": None,
        "location_quality": cell("location_quality"),
        "building_class": cell("building_class"),
        "condition": cell("condition"),
        "delivery_date": cell("delivery_date"),
        "listing_date": report_date.isoformat() if report_date else None,
        "source_key": "report_archive",
        "report_id": report_id,
        "source_url": src or f"https://report.local/{report_id}#{col}",
    }


# ── нормалізатори ────────────────────────────────────────────────────────────

_QUOTES = " \t\"'«»„“”‚‘’,;"


def _split_address_complex(raw: str) -> tuple[str, str | None]:
    text = " ".join(raw.split())
    m = re.search(r"\bжк\b[\s\"'«„“‚,]*([^\"'»“”]+)", text, flags=re.IGNORECASE)
    complex_name = None
    if m:
        complex_name = m.group(1).strip(_QUOTES).strip()
        text = text[: m.start()].strip(" ,;")
    return text, complex_name or None


def _parse_number(value: str | None) -> float | None:
    if not value:
        return None
    m = re.search(r"\d[\d\s .,]*", value)
    if not m:
        return None
    token = m.group(0).replace(" ", "").replace(" ", "")
    if "," in token and "." in token:
        token = token.replace(".", "").replace(",", ".")
    else:
        token = token.replace(",", ".")
    try:
        return float(token)
    except ValueError:
        return None


def _parse_money_usd(value: str | None) -> float | None:
    if not value:
        return None
    token = value.replace(" ", "").replace(" ", "")
    m = re.search(r"\d[\d.,]*", token)
    if not m:
        return None
    num = m.group(0)
    if num.count(",") and num.count("."):
        num = num.replace(",", "")
    else:
        num = num.replace(",", "")
    try:
        return float(num)
    except ValueError:
        return None


def _clean_url(value: str | None) -> str | None:
    if not value:
        return None
    m = re.search(r"https?://\S+", value)
    if not m:
        return None
    return m.group(0).rstrip(").,;")


def _extract_city(text: str) -> str | None:
    m = re.search(r"\bм\.?\s*([А-ЯІЇЄҐ][а-яіїєґ'’-]+)", text)
    return m.group(1) if m else None


def _extract_report_date(text: str) -> date | None:
    """Дата оцінки = найсвіжіша дата у звіті (поряд є й стара дата-еталон 2020)."""
    cands: list[date] = []
    for m in re.finditer(r"(\d{1,2})\s+([а-яіїєґ]+)\s+(\d{4})", text.lower()):
        if m.group(2) in _UA_MONTHS:
            try:
                cands.append(date(int(m.group(3)), _UA_MONTHS[m.group(2)], int(m.group(1))))
            except ValueError:
                pass
    for m in re.finditer(r"(\d{1,2})[.](\d{1,2})[.](\d{4})", text):
        try:
            cands.append(date(int(m.group(3)), int(m.group(2)), int(m.group(1))))
        except ValueError:
            pass
    cands = [d for d in cands if 2015 <= d.year <= 2030]
    return max(cands) if cands else None


# ── конвертація .doc → .docx ─────────────────────────────────────────────────

def _gather_docs(inputs: list[Path]) -> list[Path]:
    docs: list[Path] = []
    for item in inputs:
        if item.is_dir():
            docs.extend(sorted(item.rglob("*.doc")))
        elif item.suffix.lower() == ".doc" and item.exists():
            docs.append(item)
    return docs


def _convert_all(docs: list[Path], tmp: Path) -> dict[Path, Path]:
    """Конвертує всі .doc у .docx (ascii-копії, щоб уникнути проблем з кирилицею в іменах)."""
    soffice = libreoffice_path()
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) недоступний — потрібен для .doc→.docx")
    src_dir = tmp / "src"
    out_dir = tmp / "out"
    src_dir.mkdir(parents=True, exist_ok=True)
    out_dir.mkdir(parents=True, exist_ok=True)
    profile = tmp / "lo_profile"
    profile.mkdir(parents=True, exist_ok=True)

    ascii_map: dict[str, Path] = {}
    for i, doc in enumerate(docs):
        ascii_name = f"doc_{i:04d}.doc"
        shutil.copy2(doc, src_dir / ascii_name)
        ascii_map[f"doc_{i:04d}.docx"] = doc

    # Батч-конвертація пакетами (щоб не впертись у ліміт аргументів).
    files = sorted(src_dir.glob("*.doc"))
    batch = 40
    for start in range(0, len(files), batch):
        chunk = files[start:start + batch]
        command = [
            soffice, "--headless", "--nologo", "--nofirststartwizard", "--nolockcheck",
            f"-env:UserInstallation={profile.as_uri()}",
            "--convert-to", "docx", "--outdir", str(out_dir),
            *[str(f) for f in chunk],
        ]
        subprocess.run(command, capture_output=True, text=True, timeout=600)

    result: dict[Path, Path] = {}
    for docx in out_dir.glob("*.docx"):
        orig = ascii_map.get(docx.name)
        if orig is not None:
            result[orig] = docx
    return result


def _report_id(doc: Path, index: int) -> str:
    stem = re.sub(r"[^0-9a-zA-Zа-яіїєґА-ЯІЇЄҐ_-]+", "", doc.stem)[:40]
    return stem or f"r{index:04d}"


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Імпорт аналогів зі звітів (.doc) у базу.")
    parser.add_argument("--input", type=Path, action="append", required=True, help="Тека або .doc файл (можна кілька).")
    parser.add_argument("--dry-run", action="store_true", help="Парсити без запису в базу.")
    parser.add_argument("--out", type=Path, default=None, help="Куди записати JSON-звіт покриття.")
    args = parser.parse_args(argv)

    summary = import_reports(args.input, dry_run=args.dry_run, progress=lambda m: print(m, flush=True))
    text = json.dumps(summary, ensure_ascii=False, indent=2)
    print(text)
    out = args.out or (Path("data/analog_cache") / "report_import_report.json")
    try:
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(text, encoding="utf-8")
    except Exception:  # noqa: BLE001
        pass
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
