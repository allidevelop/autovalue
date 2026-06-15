from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_BREAK, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph


PLACEHOLDER_RE = re.compile(r"\{\{([A-Za-z0-9_.-]+)\}\}")
EMU_PER_INCH = 914400


class WordTemplateError(RuntimeError):
    pass


@dataclass(frozen=True)
class ParagraphLocation:
    area: str
    index: int
    paragraph: Paragraph


def office_path(path: Path) -> str:
    resolved = path.resolve()
    try:
        import win32api

        return win32api.GetShortPathName(str(resolved))
    except Exception:
        return str(resolved)


def convert_doc_to_docx(source_path: Path, output_path: Path) -> Path:
    source = source_path.resolve()
    output = output_path.resolve()
    output.parent.mkdir(parents=True, exist_ok=True)
    if source.suffix.lower() == ".docx":
        if source != output:
            shutil.copy2(source, output)
        return output
    if source.suffix.lower() != ".doc":
        raise WordTemplateError(f"Unsupported report template format: {source.suffix}")

    try:
        import pythoncom
        import win32com.client
    except Exception as exc:  # pragma: no cover - Windows dependency
        raise WordTemplateError(f"Word COM is unavailable: {exc}") from exc

    pythoncom.CoInitialize()
    word = win32com.client.DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        if _word_major_version(word) < 12:
            raise WordTemplateError(
                f"Installed Word version {word.Version} cannot reliably save .doc files as .docx. "
                "Open the source in a newer Word/LibreOffice manually and save it as .docx."
            )
        doc = word.Documents.Open(
            FileName=office_path(source),
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            PasswordDocument="",
            PasswordTemplate="",
            Revert=False,
            WritePasswordDocument="",
            WritePasswordTemplate="",
            Format=0,
            Encoding=0,
            Visible=False,
            OpenAndRepair=True,
        )
        _save_as_docx(doc, output)
        doc.Close(False)
    except WordTemplateError:
        raise
    except Exception as exc:
        raise WordTemplateError(
            "Could not convert .doc report template to .docx. "
            "Open the source in Word/LibreOffice manually, save as .docx, then run the workflow with that .docx template. "
            f"Source: {source}. Original error: {exc}"
        ) from exc
    finally:
        word.Quit()
        pythoncom.CoUninitialize()
    return output


def _word_major_version(word: Any) -> int:
    try:
        return int(str(word.Version).split(".")[0])
    except Exception:
        return 0


def _save_as_docx(doc: Any, output: Path) -> None:
    last_error: Exception | None = None
    for file_format in (16, 12):
        for method_name in ("SaveAs2", "SaveAs"):
            try:
                getattr(doc, method_name)(FileName=office_path(output), FileFormat=file_format)
                return
            except Exception as exc:
                last_error = exc
    if last_error:
        raise last_error
    raise WordTemplateError("Could not save Word document as .docx")


def load_docx_template(path: Path) -> DocumentObject:
    if path.suffix.lower() != ".docx":
        raise WordTemplateError(
            f"Word report generation requires a .docx template, got {path.suffix}. "
            "Run prepare_report_template first or save the report template as .docx."
        )
    return Document(str(path))


def iter_paragraph_locations(document: DocumentObject) -> Iterable[ParagraphLocation]:
    idx = 0
    for paragraph in document.paragraphs:
        idx += 1
        yield ParagraphLocation("body", idx, paragraph)
    for table in document.tables:
        for paragraph in _iter_table_paragraphs(table):
            idx += 1
            yield ParagraphLocation("body_table", idx, paragraph)
    for section_index, section in enumerate(document.sections, start=1):
        for part_name, part in (("header", section.header), ("footer", section.footer)):
            for paragraph in part.paragraphs:
                idx += 1
                yield ParagraphLocation(f"{part_name}_{section_index}", idx, paragraph)
            for table in part.tables:
                for paragraph in _iter_table_paragraphs(table):
                    idx += 1
                    yield ParagraphLocation(f"{part_name}_table_{section_index}", idx, paragraph)


def find_placeholders(document: DocumentObject) -> set[str]:
    placeholders: set[str] = set()
    for location in iter_paragraph_locations(document):
        placeholders.update(PLACEHOLDER_RE.findall(location.paragraph.text))
    return placeholders


def replace_text_placeholders(document: DocumentObject, values: dict[str, Any], *, skip: set[str] | None = None) -> set[str]:
    skipped = skip or set()
    replaced: set[str] = set()
    for location in iter_paragraph_locations(document):
        paragraph = location.paragraph
        if "{{" not in paragraph.text:
            continue
        replaced_in_runs = False
        for run in paragraph.runs:
            if "{{" not in run.text:
                continue
            new_run_text, run_replaced = _replace_placeholders_in_text(run.text, values, skipped)
            if new_run_text != run.text:
                run.text = new_run_text
            if run_replaced:
                replaced.update(run_replaced)
                replaced_in_runs = True
        if replaced_in_runs or "{{" not in paragraph.text:
            continue

        text = paragraph.text
        new_text, paragraph_replaced = _replace_placeholders_in_text(text, values, skipped)
        if new_text != text:
            set_paragraph_text(paragraph, new_text)
        replaced.update(paragraph_replaced)
    return replaced


def _replace_placeholders_in_text(text: str, values: dict[str, Any], skipped: set[str]) -> tuple[str, set[str]]:
    replaced: set[str] = set()

    def repl(match: re.Match[str]) -> str:
        key = match.group(1)
        if key in skipped:
            return match.group(0)
        if key not in values or values[key] is None:
            return ""
        replaced.add(key)
        return str(values[key])

    return PLACEHOLDER_RE.sub(repl, text), replaced


def replace_block_placeholder(document: DocumentObject, marker: str, callback) -> bool:
    token = f"{{{{{marker}}}}}"
    for location in iter_paragraph_locations(document):
        paragraph = location.paragraph
        if token not in paragraph.text:
            continue
        set_paragraph_text(paragraph, paragraph.text.replace(token, "").strip())
        callback(document, paragraph)
        return True
    return False


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def add_hyperlink(paragraph: Paragraph, url: str, text: str | None = None) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)

    text_node = OxmlElement("w:t")
    text_node.text = text or url
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def insert_table_after(document: DocumentObject, paragraph: Paragraph, rows: int, cols: int, style: str | None = "Table Grid") -> Table:
    table = document.add_table(rows=rows, cols=cols)
    if style:
        table.style = style
    paragraph._p.addnext(table._tbl)
    return table


def insert_page_break_after(paragraph: Paragraph) -> Paragraph:
    new_paragraph = insert_paragraph_after(paragraph)
    new_paragraph.add_run().add_break(WD_BREAK.PAGE)
    return new_paragraph


def add_picture_paragraph_after(
    anchor: Paragraph,
    image_path: Path,
    *,
    width_inches: float = 6.2,
    caption: str | None = None,
) -> Paragraph:
    caption_paragraph = anchor
    if caption:
        caption_paragraph = insert_paragraph_after(anchor, caption)
        caption_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in caption_paragraph.runs:
            run.font.size = Pt(9)
            run.font.italic = True

    image_paragraph = insert_paragraph_after(caption_paragraph)
    image_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    image_paragraph.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return image_paragraph


def remove_review_highlights(document: DocumentObject) -> int:
    """Remove yellow reviewer markup that may remain in copied Word templates."""
    removed = 0
    yellow_fills = {"FFFF00", "FFFF99", "FFF2CC", "FFF200", "FFEB3B", "FFFF66", "YELLOW"}
    roots = list(_iter_story_roots(document))
    seen_roots: set[int] = set()
    for root in roots:
        root_id = id(root)
        if root_id in seen_roots:
            continue
        seen_roots.add(root_id)
        for parent in root.iter():
            for child in list(parent):
                if child.tag == qn("w:highlight"):
                    parent.remove(child)
                    removed += 1
                    continue
                if child.tag == qn("w:shd"):
                    fill = str(child.get(qn("w:fill")) or "").upper()
                    color = str(child.get(qn("w:color")) or "").upper()
                    if fill in yellow_fills or color in yellow_fills:
                        parent.remove(child)
                        removed += 1
    return removed


def normalize_inline_image_sizes(
    document: DocumentObject,
    *,
    max_width_inches: float = 6.2,
    max_height_inches: float = 8.0,
) -> int:
    resized = 0
    max_width = int(max_width_inches * EMU_PER_INCH)
    max_height = int(max_height_inches * EMU_PER_INCH)
    for shape in document.inline_shapes:
        width = int(shape.width)
        height = int(shape.height)
        if width <= 0 or height <= 0:
            continue
        scale = min(max_width / width, max_height / height, 1.0)
        if scale >= 1.0:
            continue
        shape.width = int(width * scale)
        shape.height = int(height * scale)
        resized += 1
    return resized


def apply_basic_business_styles(document: DocumentObject) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 2" else 8)
        style.paragraph_format.space_after = Pt(6)


def _iter_table_paragraphs(table: Table) -> Iterable[Paragraph]:
    for row in table.rows:
        for cell in row.cells:
            yield from _iter_cell_paragraphs(cell)


def _iter_cell_paragraphs(cell: _Cell) -> Iterable[Paragraph]:
    for paragraph in cell.paragraphs:
        yield paragraph
    for table in cell.tables:
        yield from _iter_table_paragraphs(table)


def _iter_story_roots(document: DocumentObject) -> Iterable[Any]:
    yield document.element.body
    for section in document.sections:
        for attr in (
            "header",
            "footer",
            "first_page_header",
            "first_page_footer",
            "even_page_header",
            "even_page_footer",
        ):
            part = getattr(section, attr, None)
            element = getattr(part, "_element", None)
            if element is not None:
                yield element
