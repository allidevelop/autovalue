from __future__ import annotations

import argparse
import json
import re
import zipfile
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from pydantic import ValidationError
from rich.console import Console

from realtify.excel_tools import ExcelApp, excel_path
from realtify.intake import IntakeResult
from realtify.models import Comparable
from realtify.paths import PROJECT_ROOT, ensure_output_dir


class ReportValidationError(RuntimeError):
    pass


@dataclass(frozen=True)
class ValidationIssue:
    severity: str
    code: str
    message: str
    location: str | None = None


@dataclass(frozen=True)
class ReportValidationResult:
    ok: bool
    error_count: int
    warning_count: int
    validation_json: Path
    validation_md: Path
    issues: list[ValidationIssue]


@dataclass(frozen=True)
class ExpectedTarget:
    apartment: str | None
    total_area_m2: float | None
    total_area_comma: str
    total_area_dot: str
    address: str | None


def validate_report(
    *,
    word_path: Path,
    excel_path_value: Path | None = None,
    intake_json: Path | None = None,
    task_path: Path | None = None,
    candidates_json: Path | None = None,
    output_dir: Path | None = None,
    required_count: int = 5,
) -> ReportValidationResult:
    word = _resolve_path(word_path)
    if not word.exists():
        raise ReportValidationError(f"Word report not found: {word}")

    excel = _resolve_optional(excel_path_value)
    intake = _load_intake(_resolve_optional(intake_json))
    task = _load_yaml(_resolve_optional(task_path))
    candidates = _load_candidates(_resolve_optional(candidates_json))
    out_dir = _resolve_path(output_dir) if output_dir else word.parent
    out_dir.mkdir(parents=True, exist_ok=True)

    issues: list[ValidationIssue] = []
    expected = _expected_target(intake=intake, task=task)

    document = Document(str(word))
    visible_text = _document_visible_text(document)
    docx_xml_text = _docx_xml_text(word)
    docx_links = _docx_external_links(word)
    docx_media = _docx_media_files(word)

    _validate_placeholders(docx_xml_text, visible_text, issues)
    _validate_review_formatting(docx_xml_text, issues)
    _validate_target_text(visible_text, expected, issues)
    _validate_candidates(candidates, required_count, issues)
    _validate_candidate_links(candidates, visible_text, docx_links, issues)
    _validate_screenshots(candidates, issues)
    _validate_docx_media(candidates, docx_media, issues)
    _validate_inline_image_layout(document, issues)
    _validate_comparables_table(document, candidates, expected, issues)

    if excel:
        if not excel.exists():
            _add_error(issues, "excel_missing", f"Excel workbook not found: {excel}", str(excel))
        else:
            excel_rows = _read_adjustment_rows_from_excel(excel)
            excel_text = _excel_visible_text(excel)
            _validate_excel_target(excel_text, expected, issues)
            _validate_adjustment_table(document, excel_rows, issues)

    error_count = sum(1 for issue in issues if issue.severity == "error")
    warning_count = sum(1 for issue in issues if issue.severity == "warning")
    validation_json = out_dir / "validation.json"
    validation_md = out_dir / "validation.md"
    result = ReportValidationResult(
        ok=error_count == 0,
        error_count=error_count,
        warning_count=warning_count,
        validation_json=validation_json,
        validation_md=validation_md,
        issues=issues,
    )
    _write_validation_files(
        result,
        word=word,
        excel=excel,
        expected=expected,
        candidates_count=len(candidates),
        media_count=len(docx_media),
    )
    return result


def summarize_validation_failure(result: ReportValidationResult, *, max_issues: int = 5) -> str:
    errors = [issue for issue in result.issues if issue.severity == "error"]
    selected = errors[:max_issues]
    suffix = "" if len(errors) <= max_issues else f"; and {len(errors) - max_issues} more"
    details = "; ".join(f"{issue.code}: {issue.message}" for issue in selected)
    return f"Report validation failed: {details}{suffix}. See {result.validation_md}"


def _validate_placeholders(docx_xml_text: str, visible_text: str, issues: list[ValidationIssue]) -> None:
    placeholders = sorted(set(re.findall(r"\{\{[^{}]{1,120}\}\}", docx_xml_text + "\n" + visible_text)))
    if placeholders:
        _add_error(
            issues,
            "unresolved_placeholders",
            "Unresolved template placeholders remain in the Word report: " + ", ".join(placeholders[:10]),
            "word/document.xml",
        )


def _validate_review_formatting(docx_xml_text: str, issues: list[ValidationIssue]) -> None:
    highlight_count = len(
        re.findall(r"<w:highlight\b(?!(?:[^>]*w:val=\"none\"))", docx_xml_text, flags=re.IGNORECASE)
    )
    yellow_shading_count = len(
        re.findall(
            r"<w:shd\b[^>]*(?:w:fill|w:color)=\"(?:FFFF00|FFFF99|FFF2CC|FFF200|FFEB3B|FFFF66|yellow)\"",
            docx_xml_text,
            flags=re.IGNORECASE,
        )
    )
    if highlight_count or yellow_shading_count:
        _add_error(
            issues,
            "review_highlight_remaining",
            f"Word report still contains reviewer/highlight markup: highlights={highlight_count}, yellow_shading={yellow_shading_count}.",
            "word/*.xml",
        )


def _validate_target_text(text: str, expected: ExpectedTarget, issues: list[ValidationIssue]) -> None:
    normalized = _normalize_text(text)
    if expected.apartment:
        apartment_numbers = _contextual_apartment_numbers(normalized)
        if expected.apartment not in apartment_numbers:
            _add_error(
                issues,
                "target_apartment_missing",
                f"Expected apartment/object number {expected.apartment} was not found in target contexts.",
                "Word text",
            )
        stale_numbers = sorted(number for number in apartment_numbers if number != expected.apartment)
        for number in stale_numbers:
            _add_error(
                issues,
                "stale_target_apartment",
                f"Report still contains another apartment/object number in target context: {number}; expected {expected.apartment}.",
                "Word text",
            )

    if expected.total_area_comma:
        target_area_matches = _target_area_mentions(normalized)
        stale_areas = sorted(value for value in target_area_matches if _normalize_decimal(value) != expected.total_area_dot)
        for value in stale_areas:
            _add_error(
                issues,
                "stale_target_area",
                f"Report contains target-area phrase with {value}; expected {expected.total_area_comma}.",
                "Word text",
            )
        if expected.total_area_comma not in normalized and expected.total_area_dot not in normalized:
            _add_error(
                issues,
                "target_area_missing",
                f"Expected target area {expected.total_area_comma} was not found in the Word report.",
                "Word text",
            )


def _validate_candidates(candidates: list[Comparable], required_count: int, issues: list[ValidationIssue]) -> None:
    if len(candidates) < required_count:
        _add_error(
            issues,
            "not_enough_candidates",
            f"Selected candidates: {len(candidates)}; required: {required_count}.",
            "candidates.json",
        )
    for index, candidate in enumerate(candidates[:required_count], start=1):
        missing = []
        if candidate.area_m2 is None:
            missing.append("area_m2")
        if candidate.price_usd is None:
            missing.append("price_usd")
        if candidate.price_per_m2_usd is None:
            missing.append("price_per_m2_usd")
        if missing:
            _add_error(
                issues,
                "candidate_incomplete",
                f"Candidate {index} is missing required values: {', '.join(missing)}.",
                f"candidates[{index}]",
            )


def _validate_candidate_links(
    candidates: list[Comparable],
    visible_text: str,
    external_links: set[str],
    issues: list[ValidationIssue],
) -> None:
    haystack = visible_text + "\n" + "\n".join(external_links)
    for index, candidate in enumerate(candidates[:5], start=1):
        url = str(candidate.source_url)
        if not _contains_url(haystack, url):
            _add_error(
                issues,
                "candidate_link_missing",
                f"Candidate {index} URL is absent from visible text and hyperlink relationships: {url}",
                f"candidate {index}",
            )


def _validate_screenshots(candidates: list[Comparable], issues: list[ValidationIssue]) -> None:
    for index, candidate in enumerate(candidates[:5], start=1):
        if not candidate.screenshot_path:
            _add_error(issues, "candidate_screenshot_missing", f"Candidate {index} has no screenshot path.", f"candidate {index}")
            continue
        screenshot = Path(candidate.screenshot_path)
        if not screenshot.exists():
            _add_error(issues, "candidate_screenshot_missing", f"Candidate {index} screenshot file does not exist: {screenshot}", str(screenshot))
            continue
        try:
            from PIL import Image

            with Image.open(screenshot) as image:
                width, height = image.size
            if width < 700 or height < 900:
                _add_warning(
                    issues,
                    "candidate_screenshot_small",
                    f"Candidate {index} screenshot looks small ({width}x{height}); expected a readable full-page archive screenshot.",
                    str(screenshot),
                )
        except Exception as exc:
            _add_warning(issues, "candidate_screenshot_unchecked", f"Could not inspect screenshot {screenshot}: {exc}", str(screenshot))


def _validate_docx_media(candidates: list[Comparable], media_files: list[str], issues: list[ValidationIssue]) -> None:
    expected = min(5, len(candidates))
    if expected and len(media_files) < expected:
        _add_error(
            issues,
            "report_images_missing",
            f"Word report contains {len(media_files)} embedded media files; expected at least {expected} listing screenshots.",
            "word/media",
        )


def _validate_inline_image_layout(document, issues: list[ValidationIssue]) -> None:
    max_width = 6.35
    max_height = 8.2
    for index, shape in enumerate(document.inline_shapes, start=1):
        width = int(shape.width) / 914400
        height = int(shape.height) / 914400
        if width > max_width or height > max_height:
            _add_warning(
                issues,
                "oversized_inline_image",
                f"Inline image {index} is large for stable Word pagination ({width:.2f}x{height:.2f} inches).",
                f"inline_shapes[{index}]",
            )


def _validate_comparables_table(
    document,
    candidates: list[Comparable],
    expected: ExpectedTarget,
    issues: list[ValidationIssue],
) -> None:
    table = _find_comparables_table(document)
    if table is None:
        _add_error(issues, "comparables_table_missing", "Could not find the Word comparables table.", "Word tables")
        return

    if expected.total_area_comma:
        actual_area = _cell_text(table, 2, 6)
        if _normalize_cell(actual_area) != _normalize_cell(expected.total_area_comma):
            _add_error(
                issues,
                "comparables_target_area_mismatch",
                f"Comparables table target area is '{actual_area}', expected '{expected.total_area_comma}'.",
                "comparables table row 3 col 7",
            )

    for index, candidate in enumerate(candidates[:5], start=1):
        col = index
        expected_area = _format_decimal_comma(candidate.area_m2)
        expected_price = _format_money(candidate.price_usd)
        expected_price_m2 = _format_money(candidate.price_per_m2_usd)
        checks = [
            (2, expected_area, "candidate_area_mismatch"),
            (4, expected_price, "candidate_price_mismatch"),
            (5, expected_price_m2, "candidate_price_m2_mismatch"),
        ]
        for row, expected_value, code in checks:
            if not expected_value:
                continue
            actual = _cell_text(table, row, col)
            if _normalize_cell(actual) != _normalize_cell(expected_value):
                _add_error(
                    issues,
                    code,
                    f"Comparable {index} table value is '{actual}', expected '{expected_value}'.",
                    f"comparables table row {row + 1} col {col + 1}",
                )

        actual_url = _cell_text(table, 10, col)
        if not _contains_url(actual_url, str(candidate.source_url)):
            _add_error(
                issues,
                "comparables_source_url_mismatch",
                f"Comparable {index} source cell is '{actual_url}', expected URL {candidate.source_url}.",
                f"comparables table row 11 col {col + 1}",
            )


def _validate_excel_target(text: str, expected: ExpectedTarget, issues: list[ValidationIssue]) -> None:
    normalized = _normalize_text(text)
    if expected.apartment and expected.apartment not in normalized:
        _add_warning(
            issues,
            "excel_target_apartment_not_visible",
            f"Expected apartment/object number {expected.apartment} was not found in visible Excel cell text.",
            "Excel workbook",
        )
    if expected.total_area_comma and expected.total_area_comma not in normalized and expected.total_area_dot not in normalized:
        _add_warning(
            issues,
            "excel_target_area_not_visible",
            f"Expected target area {expected.total_area_comma} was not found in visible Excel cell text.",
            "Excel workbook",
        )


def _validate_adjustment_table(
    document,
    excel_rows: dict[int, list[str]],
    issues: list[ValidationIssue],
) -> None:
    table = _find_adjustment_table(document)
    if table is None:
        _add_error(issues, "adjustment_table_missing", "Could not find the Word adjustment/calculation table.", "Word tables")
        return
    if not excel_rows:
        _add_error(issues, "excel_adjustment_rows_missing", "Could not read adjustment rows from Excel.", "Excel workbook")
        return

    word_to_excel_rows = [
        15,
        16,
        17,
        18,
        19,
        20,
        23,
        24,
        25,
        26,
        27,
        28,
        29,
        30,
        31,
        32,
        33,
        34,
        35,
        36,
        37,
        38,
        39,
        40,
        41,
        42,
        43,
    ]
    mismatch_count = 0
    for word_row_index, excel_row_index in enumerate(word_to_excel_rows):
        expected_values = excel_rows.get(excel_row_index, [])
        seen_cells: set[int] = set()
        for col_index, expected_value in enumerate(expected_values[:8]):
            cell = table.cell(word_row_index, col_index)
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                continue
            seen_cells.add(cell_id)
            actual = _normalize_cell(cell.text)
            expected = _normalize_cell(expected_value)
            if actual == expected:
                continue
            if not actual and not expected:
                continue
            mismatch_count += 1
            if mismatch_count <= 20:
                _add_error(
                    issues,
                    "adjustment_table_excel_mismatch",
                    f"Word value '{cell.text}' differs from Excel row {excel_row_index} value '{expected_value}'.",
                    f"adjustment table row {word_row_index + 1} col {col_index + 1}",
                )
    if mismatch_count > 20:
        _add_error(
            issues,
            "adjustment_table_many_mismatches",
            f"Adjustment table has {mismatch_count} mismatched cells compared with Excel.",
            "adjustment table",
        )


def _find_comparables_table(document):
    for table in document.tables:
        if len(table.rows) < 11 or len(table.columns) < 7:
            continue
        header = _normalize_text(" ".join(cell.text for cell in table.rows[0].cells[:7]))
        row2 = _normalize_text(table.rows[2].cells[0].text).casefold()
        row10 = _normalize_text(table.rows[10].cells[0].text).casefold()
        if "№1" in header and "№5" in header and "(кв.м)" in row2 and ("джерело" in row10 or "source" in row10):
            return table
    return None


def _find_adjustment_table(document):
    for table in document.tables:
        if len(table.rows) < 27 or len(table.columns) < 8:
            continue
        header = _normalize_text(" ".join(cell.text for cell in table.rows[0].cells[:8]))
        row1 = _normalize_text(table.rows[1].cells[0].text).casefold()
        row2 = _normalize_text(table.rows[2].cells[0].text).casefold()
        if "№1" in header and "№5" in header and "1 м2" in row1 and "місцезнаходження" in row2:
            return table
    return None


def _read_adjustment_rows_from_excel(path: Path) -> dict[int, list[str]]:
    rows: dict[int, list[str]] = {}
    with ExcelApp() as excel:
        wb = excel.Workbooks.Open(excel_path(path), 0, True)
        try:
            ws = _find_calculation_sheet(wb)
            for row_index in range(15, 44):
                rows[row_index] = [str(ws.Cells(row_index, col_index).Text) for col_index in range(1, 9)]
        finally:
            wb.Close(False)
    return rows


def _excel_visible_text(path: Path) -> str:
    values: list[str] = []
    with ExcelApp() as excel:
        wb = excel.Workbooks.Open(excel_path(path), 0, True)
        try:
            for sheet_index in range(1, wb.Worksheets.Count + 1):
                ws = wb.Worksheets(sheet_index)
                used = ws.UsedRange
                start_row = int(used.Row)
                start_col = int(used.Column)
                for row_offset in range(int(used.Rows.Count)):
                    for col_offset in range(int(used.Columns.Count)):
                        text = str(ws.Cells(start_row + row_offset, start_col + col_offset).Text)
                        if text:
                            values.append(text)
        finally:
            wb.Close(False)
    return "\n".join(values)


def _find_calculation_sheet(workbook):
    for index in range(1, workbook.Worksheets.Count + 1):
        sheet = workbook.Worksheets(index)
        if "Розрах" in str(sheet.Name):
            return sheet
    return workbook.Worksheets(1)


def _expected_target(*, intake: IntakeResult | None, task: dict[str, Any]) -> ExpectedTarget:
    selected_extract = intake.selected_extract if intake else None
    selected_technical = intake.selected_technical_passport if intake else None
    target = task.get("target") if isinstance(task.get("target"), dict) else {}

    address = _first_not_empty(
        target.get("address"),
        selected_extract.address_full if selected_extract else None,
        selected_technical.address_full if selected_technical else None,
    )
    apartment = _first_not_empty(
        selected_extract.apartment_number if selected_extract else None,
        selected_technical.apartment_number if selected_technical else None,
        _parse_apartment_number(address),
    )
    total_area = _first_not_empty(
        target.get("area_m2"),
        selected_extract.total_area_m2 if selected_extract else None,
        selected_technical.total_area_m2 if selected_technical else None,
    )
    total_area_float = _to_float(total_area)
    return ExpectedTarget(
        apartment=str(apartment) if apartment not in (None, "") else None,
        total_area_m2=total_area_float,
        total_area_comma=_format_decimal_comma(total_area_float),
        total_area_dot=_format_decimal_dot(total_area_float),
        address=str(address) if address not in (None, "") else None,
    )


def _document_visible_text(document) -> str:
    values: list[str] = []
    values.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    values.append(cell.text)
    for section in document.sections:
        for part in (section.header, section.footer):
            values.extend(paragraph.text for paragraph in part.paragraphs if paragraph.text)
            for table in part.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            values.append(cell.text)
    return "\n".join(values)


def _docx_xml_text(path: Path) -> str:
    values: list[str] = []
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                values.append(archive.read(name).decode("utf-8", errors="ignore"))
    return "\n".join(values)


def _docx_external_links(path: Path) -> set[str]:
    links: set[str] = set()
    with zipfile.ZipFile(path) as archive:
        for name in archive.namelist():
            if not name.startswith("word/_rels/") or not name.endswith(".rels"):
                continue
            xml = archive.read(name).decode("utf-8", errors="ignore")
            links.update(re.findall(r'Target="([^"]+)"[^>]+TargetMode="External"', xml))
    return links


def _docx_media_files(path: Path) -> list[str]:
    with zipfile.ZipFile(path) as archive:
        return [name for name in archive.namelist() if name.startswith("word/media/")]


def _contextual_apartment_numbers(text: str) -> set[str]:
    patterns = [
        r"\b(?:квартира|квартири|квартиру|квартирою)\s*(?:№|N|No)?\s*([0-9]{1,5})\b",
        r"\bкв\.?\s*(?:№|N|No)?\s*([0-9]{1,5})\b",
    ]
    numbers: set[str] = set()
    for pattern in patterns:
        numbers.update(match.group(1) for match in re.finditer(pattern, text, flags=re.IGNORECASE))
    return numbers


def _target_area_mentions(text: str) -> set[str]:
    patterns = [
        r"\bквартир[аиюою]*\s*(?:№|N|No)?\s*[0-9]{1,5}[^.\n]{0,160}?загальн\w*\s+площею\s*([0-9]+[,.][0-9]+)",
        r"\bквартир[аиюою]*\s*(?:№|N|No)?\s*[0-9]{1,5}[^.\n]{0,160}?площею\s*([0-9]+[,.][0-9]+)\s*кв",
    ]
    values: set[str] = set()
    for pattern in patterns:
        values.update(match.group(1) for match in re.finditer(pattern, text, flags=re.IGNORECASE))
    return values


def _load_intake(path: Path | None) -> IntakeResult | None:
    if not path:
        return None
    with path.open("r", encoding="utf-8") as f:
        return IntakeResult.model_validate(json.load(f))


def _load_candidates(path: Path | None) -> list[Comparable]:
    if not path:
        return []
    with path.open("r", encoding="utf-8") as f:
        payload = json.load(f)
    raw_candidates = payload.get("candidates", [])
    try:
        return [Comparable.model_validate(item) for item in raw_candidates]
    except ValidationError as exc:
        raise ReportValidationError(f"Invalid candidates JSON: {exc}") from exc


def _load_yaml(path: Path | None) -> dict[str, Any]:
    if not path:
        return {}
    with path.open("r", encoding="utf-8") as f:
        payload = yaml.safe_load(f) or {}
    if not isinstance(payload, dict):
        raise ReportValidationError(f"{path} must contain a YAML object")
    return payload


def _write_validation_files(
    result: ReportValidationResult,
    *,
    word: Path,
    excel: Path | None,
    expected: ExpectedTarget,
    candidates_count: int,
    media_count: int,
) -> None:
    payload = {
        "ok": result.ok,
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "word_path": str(word),
        "excel_path": str(excel) if excel else None,
        "expected_target": asdict(expected),
        "candidates_count": candidates_count,
        "docx_media_count": media_count,
        "error_count": result.error_count,
        "warning_count": result.warning_count,
        "issues": [asdict(issue) for issue in result.issues],
    }
    result.validation_json.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")

    lines = [
        "# Report Validation",
        "",
        f"Created: {payload['created_at']}",
        f"Status: {'PASS' if result.ok else 'FAIL'}",
        f"Word: {word}",
        f"Excel: {excel if excel else 'not checked'}",
        f"Expected apartment/object: {expected.apartment or 'not found'}",
        f"Expected area: {expected.total_area_comma or 'not found'}",
        f"Candidates: {candidates_count}",
        f"Embedded media files: {media_count}",
        f"Errors: {result.error_count}",
        f"Warnings: {result.warning_count}",
        "",
    ]
    if result.issues:
        lines.extend(["## Issues", ""])
        for issue in result.issues:
            location = f" ({issue.location})" if issue.location else ""
            lines.append(f"- {issue.severity.upper()} `{issue.code}`{location}: {issue.message}")
    else:
        lines.append("No issues found.")
    result.validation_md.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


def _cell_text(table, row: int, col: int) -> str:
    try:
        return table.cell(row, col).text
    except IndexError:
        return ""


def _contains_url(text: str, expected_url: str) -> bool:
    normalized_text = text.replace("&amp;", "&")
    variants = {expected_url, expected_url.rstrip("/")}
    return any(variant and variant in normalized_text for variant in variants)


def _normalize_text(value: Any) -> str:
    return re.sub(r"[ \t\r\f\v]+", " ", str(value or "").replace("\xa0", " ")).strip()


def _normalize_cell(value: Any) -> str:
    text = _normalize_text(value)
    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def _normalize_decimal(value: Any) -> str:
    return str(value or "").replace(",", ".").strip()


def _format_decimal_comma(value: Any, *, decimals: int = 2, trim: bool = True) -> str:
    number = _to_float(value)
    if number is None:
        return ""
    text = f"{number:.{decimals}f}".replace(".", ",")
    if trim and "," in text:
        text = text.rstrip("0").rstrip(",")
    return text


def _format_decimal_dot(value: Any, *, decimals: int = 2, trim: bool = True) -> str:
    number = _to_float(value)
    if number is None:
        return ""
    text = f"{number:.{decimals}f}"
    if trim and "." in text:
        text = text.rstrip("0").rstrip(".")
    return text


def _format_money(value: Any) -> str:
    number = _to_float(value)
    if number is None:
        return ""
    return f"{number:,.0f}".replace(",", " ")


def _to_float(value: Any) -> float | None:
    if value in (None, ""):
        return None
    try:
        return float(str(value).replace(",", "."))
    except (TypeError, ValueError):
        return None


def _parse_apartment_number(text: Any) -> str | None:
    match = re.search(r"(?:квартира|кв\.?)\s*(?:№|N|No)?\s*([0-9]{1,5})", str(text or ""), re.IGNORECASE)
    return match.group(1) if match else None


def _first_not_empty(*values: Any) -> Any:
    for value in values:
        if value not in (None, ""):
            return value
    return None


def _add_error(issues: list[ValidationIssue], code: str, message: str, location: str | None = None) -> None:
    issues.append(ValidationIssue(severity="error", code=code, message=message, location=location))


def _add_warning(issues: list[ValidationIssue], code: str, message: str, location: str | None = None) -> None:
    issues.append(ValidationIssue(severity="warning", code=code, message=message, location=location))


def _resolve_optional(path: Path | None) -> Path | None:
    if not path:
        return None
    return _resolve_path(path)


def _resolve_path(path: Path) -> Path:
    return path if path.is_absolute() else PROJECT_ROOT / path


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Validate generated Word valuation report against intake, candidates, and Excel.")
    parser.add_argument("--word", type=Path, required=True, help="Generated .docx valuation report.")
    parser.add_argument("--excel", type=Path, default=None, help="Filled Excel workbook.")
    parser.add_argument("--intake", type=Path, default=None, help="intake.json from PDF extraction.")
    parser.add_argument("--task", type=Path, default=None, help="task.generated.yaml from PDF extraction.")
    parser.add_argument("--candidates", type=Path, default=None, help="candidates.json selected for report.")
    parser.add_argument("--out", type=Path, default=None, help="Directory for validation.json and validation.md.")
    parser.add_argument("--required-count", type=int, default=5)
    args = parser.parse_args(argv)

    console = Console()
    try:
        result = validate_report(
            word_path=args.word,
            excel_path_value=args.excel,
            intake_json=args.intake,
            task_path=args.task,
            candidates_json=args.candidates,
            output_dir=args.out,
            required_count=args.required_count,
        )
    except Exception as exc:
        console.print(f"[red]Report validation failed to run:[/red] {exc}")
        return 1

    if result.ok:
        console.print("[green]Report validation passed[/green]")
    else:
        console.print(f"[red]Report validation failed[/red]: {result.error_count} errors, {result.warning_count} warnings")
        for issue in result.issues[:20]:
            style = "red" if issue.severity == "error" else "yellow"
            console.print(f"[{style}]{issue.severity.upper()}[/{style}] {issue.code}: {issue.message}")
    console.print(f"Validation JSON: {result.validation_json}")
    console.print(f"Validation report: {result.validation_md}")
    return 0 if result.ok else 1


if __name__ == "__main__":
    raise SystemExit(main())
