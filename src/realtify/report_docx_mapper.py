"""Рендер структурованого документа звіту → .docx (python-docx) з нуля за тією ж
схемою нод і тією ж style-spec, що й HTML-рендерер. Жодного шаблону-.docx і media-swap:
кожна нода будується програмно, геометрія (поля, ширини колонок) — у мм зі spec.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor

from realtify import report_styles as styles
from realtify.report_styles import units

_ALIGN = {
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
    "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def _set_run_font(run, family: str, size_pt: float, *, bold=False, italic=False, underline=False, shade: str | None = None) -> None:
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), family)
    if rfonts.getparent() is None:
        rpr.insert(0, rfonts)
    if shade:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), shade)
        rpr.append(shd)


def _set_table_borders(table, color: str, size_pt: float) -> None:
    sz = max(2, round(size_pt * 8))  # w:sz у 1/8 pt
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def _set_cell_margins(table, v_mm: float, h_mm: float) -> None:
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for edge, mm in (("top", v_mm), ("bottom", v_mm), ("left", h_mm), ("right", h_mm)):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(round(mm * 56.6929)))  # mm → twips
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def _fixed_layout(table) -> None:
    table.allow_autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)


def _configure_section(doc, spec: dict[str, Any]) -> None:
    page, m = spec["page"], spec["page"]["marginMm"]
    s = doc.sections[0]
    s.page_width, s.page_height = Mm(page["widthMm"]), Mm(page["heightMm"])
    s.top_margin, s.bottom_margin = Mm(m["top"]), Mm(m["bottom"])
    s.left_margin, s.right_margin = Mm(m["left"]), Mm(m["right"])


def _emit_inline(paragraph, nodes: list[dict], spec: dict[str, Any], size_pt: float, mode: str, *, force_bold=False) -> None:
    fam = spec["fonts"]["body"]["family"]
    show_prov = styles.export_mode(spec, mode).get("showProvenance", False)
    for n in nodes or []:
        t = n.get("type")
        if t == "text":
            marks = {m.get("type") for m in n.get("marks", []) or []}
            run = paragraph.add_run(n.get("text", ""))
            _set_run_font(run, fam, size_pt, bold=force_bold or "bold" in marks,
                          italic="italic" in marks, underline="underline" in marks)
        elif t == "variableField":
            a = n.get("attrs") or {}
            shade = styles.provenance_color(spec, a.get("source", "auto")).get("wordShade") if show_prov else None
            run = paragraph.add_run(a.get("value", ""))
            _set_run_font(run, fam, size_pt, bold=force_bold, shade=shade)


def _emit_table(doc, node: dict, spec: dict[str, Any]) -> None:
    a = node.get("attrs") or {}
    kind = a.get("kind", "")
    ts = styles.table_style(spec, kind)
    cols = a.get("columnsMm") or ts.get("columnsMm", [])
    header = a.get("header") or []
    rows = a.get("rows") or []
    fam = spec["fonts"]["body"]["family"]
    size_pt = ts.get("fontSizePt", 8)
    n_cols = len(cols) or (len(header) if header else (len(rows[0]) if rows else 1))
    table = doc.add_table(rows=0, cols=n_cols)
    _fixed_layout(table)
    _set_table_borders(table, ts.get("borderColor", "000000"), ts.get("borderPt", 0.5))
    _set_cell_margins(table, ts.get("cellPadVMm", 0.3), ts.get("cellPaddingMm", 1.0))
    align = _ALIGN.get(ts.get("align", "center"), WD_ALIGN_PARAGRAPH.CENTER)

    def fill(cells_text: list[str], bold: bool) -> None:
        row = table.add_row()
        for i, val in enumerate(cells_text[:n_cols]):
            cell = row.cells[i]
            if i < len(cols):
                cell.width = Mm(cols[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            para = cell.paragraphs[0]
            para.alignment = align
            pf = para.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = Pt(size_pt + 0.5)  # точна висота рядка — тугі рядки таблиці
            run = para.add_run("" if val is None else str(val))
            _set_run_font(run, fam, size_pt, bold=bold)

    if header:
        fill(header, ts.get("headerBold", True))
    for r in rows:
        fill(list(r), False)


def _emit_image(doc, node: dict, spec: dict[str, Any]) -> None:
    import base64
    import io

    a = node.get("attrs") or {}
    src = str(a.get("srcRef", ""))
    if not src.startswith("data:") or "," not in src:
        return
    try:
        data = base64.b64decode(src.split(",", 1)[1])
    except Exception:  # noqa: BLE001
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    w = int(a.get("widthEmu") or 0)
    try:
        p.add_run().add_picture(io.BytesIO(data), width=Emu(w) if w else None)
    except Exception:  # noqa: BLE001
        return
    cap, href = a.get("caption"), a.get("href")
    if cap or href:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        size = styles.block_style(spec, "caption").get("sizePt", 11)
        if href:
            from realtify.word_tools import add_hyperlink
            add_hyperlink(cp, str(href), str(cap or href))
        else:
            r = cp.add_run(str(cap))
            _set_run_font(r, spec["fonts"]["body"]["family"], size)


def _emit_block(doc, node: dict, spec: dict[str, Any], mode: str) -> None:
    t = node.get("type")
    if t == "heading":
        lvl = (node.get("attrs") or {}).get("level", 2)
        num = (node.get("attrs") or {}).get("numbering")
        b = styles.block_style(spec, f"heading{lvl}")
        p = doc.add_paragraph()
        p.alignment = _ALIGN.get(b.get("align", "left"), WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_before = Pt(b.get("spaceBeforePt", 0))
        p.paragraph_format.space_after = Pt(b.get("spaceAfterPt", 0))
        if num:
            r = p.add_run(f"{num} ")
            _set_run_font(r, spec["fonts"]["body"]["family"], b.get("sizePt", 14), bold=b.get("bold", True))
        _emit_inline(p, node.get("content", []), spec, b.get("sizePt", 14), mode, force_bold=b.get("bold", True))
    elif t in ("paragraph", "definitionItem"):
        b = styles.block_style(spec, "paragraph")
        p = doc.add_paragraph()
        align = (node.get("attrs") or {}).get("align", b.get("align", "justify"))
        p.alignment = _ALIGN.get(align, WD_ALIGN_PARAGRAPH.JUSTIFY)
        pf = p.paragraph_format
        pf.line_spacing = b.get("lineHeight", 1.4)
        pf.space_after = Pt(b.get("spaceAfterPt", 0))
        if t == "definitionItem":
            term = (node.get("attrs") or {}).get("term", "")
            if term:
                r = p.add_run(term + " ")
                _set_run_font(r, spec["fonts"]["body"]["family"], b.get("sizePt", 14), bold=True)
        else:
            pf.first_line_indent = Mm(b.get("firstLineIndentMm", 0))
        _emit_inline(p, node.get("content", []), spec, b.get("sizePt", 14), mode)
    elif t == "table":
        _emit_table(doc, node, spec)
    elif t in ("image", "documentScan"):
        _emit_image(doc, node, spec)
    elif t == "pageBreak":
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    elif t in ("bulletList", "orderedList"):
        style = "List Bullet" if t == "bulletList" else "List Number"
        b = styles.block_style(spec, "paragraph")
        for li in node.get("content", []):
            p = doc.add_paragraph(style=style)
            _emit_inline(p, li.get("content", []), spec, b.get("sizePt", 14), mode)


def render_document_docx(document: dict, out_path: Path, spec: dict[str, Any] | None = None, mode: str = "clean") -> Path:
    spec = spec or styles.load_style_spec()
    doc = Document()
    _configure_section(doc, spec)
    normal = doc.styles["Normal"]
    normal.font.name = spec["fonts"]["body"]["family"]
    normal.font.size = Pt(spec["fonts"]["body"]["sizePt"])
    normal.font.color.rgb = RGBColor(0, 0, 0)
    # тугий baseline — інакше Word-овський дефолтний інтервал роздуває рядки таблиці
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    for node in document.get("content", []):
        _emit_block(doc, node, spec, mode)
    out_path = Path(out_path)
    doc.save(str(out_path))
    return out_path
