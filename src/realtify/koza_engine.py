"""Коза-движок: генерація звіту КЛОНУВАННЯМ готового звіту-шаблона («кози») того ж
будинку з бібліотеки + підстановка даних конкретної квартири.

Ідея (вимога клієнта): для будинку, по якому вже є готовий звіт у базі шаблонів,
не досліджуємо аналоги заново і не збираємо з дженерик-шаблона — беремо козу цього
будинку (за ключем дому) і клонуємо її, підставляючи дані нової квартири: № квартири,
площу, кімнатність, дати, вартість, скани витяга/техпаспорта, № витяга. Аналоги,
таблиці 7.1/7.2/7.3, опис ЖК, характеристики, формат — лишаються як у козі.

Модуль:
- build_index(reports_dir): конвертує .doc→текст, дістає адресу, ключ дому → коза(и);
- find_koza(building_key, index): підбір кози для дому;
- koza_apartment_values(text): старі значення квартири кози (для find/replace).
Сам клон-рендер (підстановка у .docx + свопи сканів) — у clone_koza().
"""
from __future__ import annotations

import json
import os
import re
import subprocess
from pathlib import Path
from typing import Any

from realtify.analog_cache import building_match_key
from realtify.paths import PROJECT_ROOT

REPORTS_LIBRARY = PROJECT_ROOT / "data" / "reports_all"
KOZA_INDEX_PATH = PROJECT_ROOT / "data" / "koza_index.json"


def _soffice_convert(path: Path, outdir: Path, convert_to: str, out_ext: str) -> Path | None:
    """Конвертація через LibreOffice з УНІКАЛЬНИМ профілем на виклик (інакше
    конкурентні soffice б'ються за один профіль-lock і мовчки падають). Перевіряємо
    наявність вихідного файлу."""
    import shutil

    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        return None
    profile = Path(outdir) / "_loprofile"
    subprocess.run(
        [soffice, "--headless", f"-env:UserInstallation={profile.as_uri()}",
         "--nolockcheck", "--nofirststartwizard", "--nologo",
         "--convert-to", convert_to, "--outdir", str(outdir), str(path)],
        capture_output=True, timeout=180,
    )
    out = Path(outdir) / (path.stem + out_ext)
    return out if out.exists() else None


def _doc_to_text(path: Path) -> str:
    """Конвертує .doc/.docx у текст через LibreOffice (unicode-safe)."""
    import tempfile
    with tempfile.TemporaryDirectory(prefix="koza_txt_") as tmp:
        out = _soffice_convert(path, Path(tmp), "txt:Text", ".txt")
        return out.read_text(encoding="utf-8", errors="ignore") if out else ""


def extract_address(text: str) -> tuple[str | None, str | None]:
    """Дістає (city, address) з тексту звіту-кози (рядок «Місце розташування:» або
    «за адресою:»). Повертає сирий адресний рядок — далі building_match_key."""
    m = (re.search(r"Місце розташування:\s*\n?\s*([^\n]+)", text)
         or re.search(r"за адресою:\s*\n?\s*([^\n]+)", text)
         or re.search(r"Адреса[^:\n]{0,30}:\s*\n?\s*([^\n]+)", text))
    if not m:
        return None, None
    addr = " ".join(m.group(1).split()).strip().rstrip(".")
    city = None
    cm = re.search(r"м\.?\s*([А-ЯІЇЄҐ][а-яіїєґ'\-]+)", addr)
    if cm:
        city = cm.group(1)
    return city, addr


def koza_apartment_values(text: str) -> dict[str, Any]:
    """Старі значення квартири кози (для пошуку-заміни при клонуванні)."""
    v: dict[str, Any] = {}
    m = re.search(r"кварти\w*\s*№?\s*(\d+)", text)
    v["apartment"] = m.group(1) if m else None
    m = re.search(r"загальною площею\s*([\d]+[.,]?\d*)\s*кв", text)
    v["area"] = m.group(1) if m else None
    m = re.search(r"(одно|дво|три|чотири|п['’]?яти|шести)кімнатн", text, re.IGNORECASE)
    v["rooms_word_stem"] = m.group(1).lower() if m else None
    m = re.search(r"Дата оцінки:\s*([^\n]+)", text)
    v["valuation_date"] = m.group(1).strip() if m else None
    m = re.search(r"Дата складання звіту:\s*([^\n]+)", text)
    v["report_date"] = m.group(1).strip() if m else None
    m = re.search(r"Ринкова вартість[^\d]*([\d\s]+[,.]\d{2})", text)
    v["market_value"] = m.group(1).strip() if m else None
    m = re.search(r"Індексний номер витягу:\s*([0-9]+)", text)
    v["extract_index_number"] = m.group(1) if m else None
    return v


def build_index(reports_dir: Path = REPORTS_LIBRARY, *, save: bool = True,
                progress=None) -> dict[str, list[str]]:
    """Будує індекс {building_key: [імена файлів кози]} з бібліотеки звітів.
    Конвертує кожен .doc у текст і дістає адресу (контент-based, без залежності
    від кодування імені файлу)."""
    index: dict[str, list[str]] = {}
    docs = sorted(reports_dir.glob("*.doc")) + sorted(reports_dir.glob("*.docx"))
    for i, f in enumerate(docs):
        if progress:
            progress(f"[{i + 1}/{len(docs)}] {f.name}")
        try:
            text = _doc_to_text(f)
            city, addr = extract_address(text)
            if not addr:
                continue
            key = building_match_key(city=city, address=addr)
            if key:
                index.setdefault(key, []).append(f.name)
        except Exception:  # noqa: BLE001 — один битий звіт не валить індекс
            continue
    if save:
        KOZA_INDEX_PATH.parent.mkdir(parents=True, exist_ok=True)
        KOZA_INDEX_PATH.write_text(json.dumps(index, ensure_ascii=False, indent=1), encoding="utf-8")
    return index


def load_index() -> dict[str, list[str]]:
    if KOZA_INDEX_PATH.exists():
        return json.loads(KOZA_INDEX_PATH.read_text(encoding="utf-8"))
    return {}


def _save_index(index: dict[str, list[str]]) -> None:
    KOZA_INDEX_PATH.parent.mkdir(parents=True, exist_ok=True)
    KOZA_INDEX_PATH.write_text(json.dumps(index, ensure_ascii=False, indent=1), encoding="utf-8")


def index_status() -> dict[str, Any]:
    """Статус бази шаблонів (козлів): к-сть домів і всього файлів."""
    index = load_index()
    total = sum(len(v) for v in index.values())
    return {"buildings": len(index), "total": total, "exists": bool(index)}


def add_koza(koza_path: Path) -> str | None:
    """Додає вже збережений у бібліотеку файл-козу в індекс (за ключем дому).
    Повертає ключ дому або None, якщо адресу не вдалося дістати."""
    text = _doc_to_text(koza_path)
    city, addr = extract_address(text)
    if not addr:
        return None
    key = building_match_key(city=city, address=addr)
    if not key:
        return None
    index = load_index()
    names = index.setdefault(key, [])
    if koza_path.name not in names:
        names.append(koza_path.name)
    _save_index(index)
    return key


def _doc_to_docx(path: Path, outdir: Path) -> Path | None:
    return _soffice_convert(path, Path(outdir), "docx:MS Word 2007 XML", ".docx")


def _iter_paragraphs(doc):
    """Усі абзаци документа: тіло + клітинки таблиць (включно з вкладеними)."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                for nt in cell.tables:
                    for nr in nt.rows:
                        for nc in nr.cells:
                            yield from nc.paragraphs


def _replace_in_paragraph(p, repls: list[tuple[str, str]]) -> bool:
    """Run-aware заміна old→new в абзаці. Редагує ЛИШЕ ті runs, що перекривають збіг
    (зберігає форматування решти), коректно обробляє збіг через кілька runs, і НЕ
    застосовує заміну повторно (пошук просувається за вставленим new — безпечно навіть
    коли new містить old, напр. 20→205)."""
    changed = False
    runs = p.runs
    if not runs:
        return False
    for old, new in repls:
        if not old:
            continue
        search_start = 0
        while True:
            texts = [r.text for r in runs]
            full = "".join(texts)
            idx = full.find(old, search_start)
            if idx < 0:
                break
            end = idx + len(old)
            # знаходимо runs, що перекривають [idx, end)
            pos = 0
            sr = so = er = eo = None
            for ri, t in enumerate(texts):
                nxt = pos + len(t)
                if sr is None and nxt > idx:
                    sr, so = ri, idx - pos
                if er is None and nxt >= end:
                    er, eo = ri, end - pos
                    break
                pos = nxt
            if sr is None or er is None:
                break
            if sr == er:
                t = texts[sr]
                runs[sr].text = t[:so] + new + t[eo:]
            else:
                runs[sr].text = texts[sr][:so] + new
                for ri in range(sr + 1, er):
                    runs[ri].text = ""
                runs[er].text = texts[er][eo:]
            changed = True
            search_start = idx + len(new)  # за вставленим new → без повторної заміни
    return changed


def _format_uk_date(d) -> str:
    months = ["січня", "лютого", "березня", "квітня", "травня", "червня",
              "липня", "серпня", "вересня", "жовтня", "листопада", "грудня"]
    return f"{d.day:02d} {months[d.month - 1]} {d.year} року"


def _set_extract_ref(doc, new_ref: str) -> None:
    """Оновлює реквізит витяга у фразі «…про реєстрацію прав власності № … від … р.»:
    new_ref непорожнє → ставимо новий № (нової квартири); порожнє → прибираємо
    (опція клієнта; справжній номер усе одно видно на вклеєному скані витяга)."""
    pat = re.compile(r"(прав власності)\s*№\s*\d[\d  ]*(?:\s*від\s+[\d.  ]+р?\.?)?", re.IGNORECASE)
    tail = f" {new_ref}" if new_ref else ""
    for p in _iter_paragraphs(doc):
        if not p.runs:
            continue
        full = "".join(r.text for r in p.runs)
        new = pat.sub(lambda m: m.group(1) + tail, full)
        if new != full:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ""


def clone_koza(koza_path: Path, out_path: Path, *, old: dict, new: dict,
               scans: dict | None = None, progress=None, extract_ref: str | None = None) -> Path:
    """Клонує козу (.doc/.docx) у out_path (.docx), підставляючи дані нової квартири.
    old/new — словники значень кози й нової квартири (apartment, area, rooms_word_stem,
    valuation_date, report_date, market_value, extract_index_number).
    scans (опц.) — {'vityag':png,'techpass':png} замінює скани кози на скани квартири."""
    import tempfile

    from docx import Document
    with tempfile.TemporaryDirectory(prefix="koza_clone_") as tmp:
        src = koza_path
        if koza_path.suffix.lower() == ".doc":
            conv = _doc_to_docx(koza_path, Path(tmp))
            if conv is None:
                raise RuntimeError("LibreOffice не сконвертував козу в .docx")
            src = conv
        doc = Document(str(src))

        repls: list[tuple[str, str]] = []
        # № квартири — ТІЛЬКИ у контексті «квартир*/кв» (щоб не зачепити реєстраційні
        # номери, № аналогів, рік тощо). Голий «№ N» НЕ використовуємо.
        if old.get("apartment") and new.get("apartment"):
            oa, na = str(old["apartment"]), str(new["apartment"])
            for pat in (f"квартири № {oa}", f"квартира № {oa}", f"квартири №{oa}", f"квартира №{oa}",
                        f"квартири {oa}", f"квартира {oa}", f"кв. {oa}", f"кв {oa}"):
                repls.append((pat, pat.replace(oa, na)))
        # площа — ТІЛЬКИ у контексті «площею {N}» (прозовий опис об'єкта). Голу заміну
        # числа НЕ робимо: інакше площа потрапляла б у площі аналогів/числа таблиць.
        # Таблиці (об'єкт-колонка) лишаються з кози до етапу точного розрахунку (L).
        if old.get("area") and new.get("area"):
            repls.append((f"площею {old['area']}", f"площею {new['area']}"))
        # дати
        for k in ("valuation_date", "report_date"):
            if old.get(k) and new.get(k):
                repls.append((str(old[k]), str(new[k])))
        # вартість
        if old.get("market_value") and new.get("market_value"):
            repls.append((str(old["market_value"]), str(new["market_value"])))
        # кімнатність: стем кози (як витягнуто, з її апострофом) → новий стем.
        # Випадок (кімнатної/кімнатна) — у закінченні, його не чіпаємо.
        os_, ns_ = old.get("rooms_word_stem"), new.get("rooms_word_stem")
        if os_ and ns_ and os_ != ns_:
            repls.append((f"{os_}кімнат", f"{ns_}кімнат"))
            repls.append((f"{os_[:1].upper()}{os_[1:]}кімнат", f"{ns_[:1].upper()}{ns_[1:]}кімнат"))

        for p in _iter_paragraphs(doc):
            _replace_in_paragraph(p, repls)

        if extract_ref is not None:
            _set_extract_ref(doc, extract_ref)

        if scans:
            try:
                n = swap_scans(doc, scans)
                if n == 0 and progress:
                    progress("Увага: скани витяга/техпаспорта в козі не розпізнані — "
                             "у клоні лишилися скани кози, перевірте/замініть вручну.")
            except Exception as exc:  # noqa: BLE001 — своп сканів не має валити клон
                if progress:
                    progress(f"Своп сканів не вдався: {exc}")

        out_path = Path(out_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
    return out_path


def _scan_drawings(doc) -> list:
    """Знаходить великі ПОРТРЕТНІ малюнки (повносторінкові скани витяга/техпаспорта)
    у порядку документа. Логотипи/підписи/фото — менші або альбомні, відсіюються."""
    from docx.oxml.ns import qn
    parts = doc.part.related_parts
    out = []
    for dr in doc.element.body.iter(qn("w:drawing")):
        blip = dr.find(".//" + qn("a:blip"))
        ext = dr.find(".//" + qn("wp:extent"))
        if blip is None or ext is None:
            continue
        rid = blip.get(qn("r:embed"))
        part = parts.get(rid)
        if part is None:
            continue
        cx, cy = int(ext.get("cx") or 0), int(ext.get("cy") or 0)
        if len(part.blob) > 800_000 and cy > cx:  # велике + портретне → скан
            out.append((dr, blip, ext))
    return out


def swap_scans(doc, scans: dict) -> int:
    """Замінює скани кози (витяг/техпаспорт) на скани нової квартири. scans —
    {'vityag': png_bytes, 'techpass': png_bytes}. Повертає к-сть замінених."""
    import io

    from docx.oxml.ns import qn
    order = ["vityag", "techpass"]
    swapped = 0
    for i, (_dr, blip, _ext) in enumerate(_scan_drawings(doc)):
        kind = order[i] if i < len(order) else None
        png = scans.get(kind) if kind else None
        if not png:
            continue
        rid, _image_part = doc.part.get_or_add_image(io.BytesIO(png))
        blip.set(qn("r:embed"), rid)
        swapped += 1
    return swapped


def output_filename(*, city: str | None, street: str | None, building: str | None,
                    apartment: str | None, ext: str = "docx") -> str:
    """Ім'я вихідного файлу у клієнтському форматі (вимога H):
    «Звіт м. {місто}, вул. {вулиця} буд. {будинок} кв {N}.docx»."""
    parts = ["Звіт"]
    if city:
        parts.append(f"м. {city},")
    if street:
        parts.append(f"вул. {street}")
    if building:
        parts.append(f"буд. {building}")
    if apartment:
        parts.append(f"кв {apartment}")
    name = " ".join(parts).strip()
    name = re.sub(r"[\\/:*?\"<>|]", "", name)
    return f"{name}.{ext}"


_STEM_BY_ROOMS = {1: "одно", 2: "дво", 3: "три", 4: "чотири", 5: "п'яти", 6: "шести"}


def _rooms_stem(rooms: Any) -> str | None:
    try:
        return _STEM_BY_ROOMS.get(int(rooms))
    except (TypeError, ValueError):
        return None


def _street_building(address: str | None) -> tuple[str | None, str | None]:
    a = address or ""
    sm = re.search(r"вул(?:иця|\.)?\s*(.+?)\s*,?\s*(?:буд(?:инок|\.)?|$)", a)
    bm = re.search(r"буд(?:инок|\.)?\s*([0-9][0-9A-Za-zА-Яа-яІЇЄҐіїєґ\-/]*)", a)
    street = sm.group(1).strip().rstrip(",") if sm else None
    building = bm.group(1).strip() if bm else None
    return street, building


def _koza_output_name(koza_path: Path, old_apt: str | None, new_apt: str | None,
                      *, fallback_stem: str | None = None) -> str:
    """Ім'я вихідного файлу у форматі клієнта. Якщо № квартири кози — ОСТАННЄ число
    в імені (клієнтський формат «…кв N»), безпечно його замінюємо. Інакше не чіпаємо
    інші числа (напр. № дому) — будуємо ім'я з адреси (fallback_stem)."""
    stem = koza_path.stem
    if new_apt and old_apt and re.search(rf"\b{re.escape(str(old_apt))}\b(?=\D*$)", stem):
        stem = re.sub(rf"\b{re.escape(str(old_apt))}\b(?=\D*$)", str(new_apt), stem, count=1)
        return f"{stem}.docx"
    if fallback_stem:
        return f"{fallback_stem}.docx"
    return f"{stem}.docx"


def build_report_via_koza(obj_dir: Path, out_dir: Path, *,
                          index: dict[str, list[str]] | None = None,
                          progress=None) -> Path | None:
    """Якщо для будинку об'єкта є коза — будує звіт КЛОНУВАННЯМ кози з підстановкою
    даних квартири (№/площа/кімнатність/дати/вартість/скани). Інакше повертає None
    (далі — звичайний пайплайн). Вартість на цьому етапі — пропорційна (ціна/м² кози
    × нова площа); точний розрахунок (торг тощо) — окремим етапом за Excel клієнта."""
    import tempfile

    from realtify.excel_summary import read_excel_report_values
    from realtify.report_generator import (
        _load_candidates, _load_intake, _load_yaml, _resolve_optional, build_report_values,
    )
    from realtify.report_scans import render_object_scans
    from realtify.valuation_date import resolve_valuation_context

    obj_dir = Path(obj_dir)
    intake = _load_intake(_resolve_optional(obj_dir / "intake.json"))
    task = _load_yaml(_resolve_optional(obj_dir / "task.generated.yaml"))
    candidates = _load_candidates(_resolve_optional(obj_dir / "candidates.json"))
    excels = sorted(obj_dir.glob("*_filled.xls")) or sorted(obj_dir.glob("*.xls"))
    excel_path = excels[0] if excels else None

    values = build_report_values(
        intake=intake, task=task, candidates=candidates,
        excel_path=excel_path, excel_values=read_excel_report_values(excel_path),
    )
    city = values.get("report_city") or None
    address = values.get("address_full") or None
    koza = find_koza(city=city, address=address, index=index)
    if not koza:
        if progress:
            progress(f"Коза для дому не знайдена ({address}) — звичайний пайплайн.")
        return None
    if progress:
        progress(f"Знайдено козу дому: {koza.name} — клонування під квартиру.")

    old = koza_apartment_values(_doc_to_text(koza))

    # Вартість: пропорційно до площі кози (ціна/м² стабільна для дому).
    def _num(s: Any) -> float | None:
        if not s:
            return None
        try:
            return float(str(s).replace("\xa0", "").replace(" ", "").replace(",", "."))
        except ValueError:
            return None

    new_area_num = _num(values.get("total_area_m2"))
    koza_val, koza_area = _num(old.get("market_value")), _num(old.get("area"))
    market_value = None
    if koza_val and koza_area and new_area_num:
        v = round(koza_val / koza_area * new_area_num, -2)
        market_value = f"{v:,.2f}".replace(",", " ").replace(".", ",")

    ctx = resolve_valuation_context(task, excel_path=excel_path)
    new = {
        "apartment": values.get("apartment_number") or None,
        "area": values.get("total_area_m2") or None,
        "rooms_word_stem": _rooms_stem(values.get("rooms_count")),
        "valuation_date": _format_uk_date(ctx.valuation_date),
        "report_date": _format_uk_date(ctx.valuation_date),
        "market_value": market_value,
        "extract_index_number": values.get("extract_index_number") or None,
    }
    if not new["apartment"]:
        # без № квартири клон неможливий (і виникають колізії імен) → звичайний пайплайн.
        if progress:
            progress("№ квартири об'єкта не визначено — коза-клон пропущено.")
        return None

    street, building = _street_building(address)
    fallback_stem = output_filename(
        city=city, street=street, building=building, apartment=new["apartment"]
    ).rsplit(".docx", 1)[0]

    scans = {}
    with tempfile.TemporaryDirectory(prefix="koza_scans_") as tmp:
        try:
            raw = render_object_scans(intake, task, Path(tmp))
            scans = {k: v[0] for k, v in raw.items() if v and v[0]}
        except Exception:  # noqa: BLE001 — скани не критичні для клону
            scans = {}
        out_dir = Path(out_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / _koza_output_name(
            koza, old.get("apartment"), new["apartment"], fallback_stem=fallback_stem)
        # № витяга: за замовч. ставимо новий реквізит (OCR 300dpi тепер точний);
        # REALTIFY_KOZA_EXTRACT_REF=drop — прибрати (скан усе одно показує номер).
        drop = os.getenv("REALTIFY_KOZA_EXTRACT_REF", "").strip().lower() in {"drop", "off", "0", "no", "remove"}
        extract_ref = "" if drop else (values.get("extract_reference") or "")
        clone_koza(koza, out_path, old=old, new=new, scans=scans, progress=progress,
                   extract_ref=extract_ref)
    # Перевіряємо, що клон відкривається (інакше не віддаємо як первинний артефакт).
    try:
        from docx import Document as _Doc
        _Doc(str(out_path))
    except Exception as exc:  # noqa: BLE001
        if progress:
            progress(f"Коза-клон не відкривається ({exc}) — фолбек на звичайний пайплайн.")
        return None
    return out_path


def find_koza(*, city: str | None, address: str | None,
              index: dict[str, list[str]] | None = None,
              reports_dir: Path = REPORTS_LIBRARY) -> Path | None:
    """Підбирає козу для будинку об'єкта за ключем дому. Повертає шлях до .doc/.docx."""
    if not address:
        return None
    index = index if index is not None else load_index()
    key = building_match_key(city=city, address=address)
    names = index.get(key)
    if not names:
        return None
    # Беремо найсвіжіший файл (за mtime) як представника будинку.
    paths = [reports_dir / n for n in names if (reports_dir / n).exists()]
    if not paths:
        return None
    return max(paths, key=lambda p: p.stat().st_mtime)
