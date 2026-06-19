from __future__ import annotations

import argparse
import json
from dataclasses import dataclass
from datetime import date, datetime
from pathlib import Path
from typing import Any

import yaml
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from pydantic import ValidationError
from rich.console import Console

from realtify.excel_summary import read_excel_report_values
from realtify.excel_sidecar import sidecar_adjustment_rows
from realtify.excel_tools import excel_path
from realtify.intake import IntakeResult
from realtify.models import Comparable
from realtify.paths import PROJECT_ROOT, ensure_output_dir
from realtify.valuation_date import resolve_valuation_date
from realtify.word_tools import (
    WordTemplateError,
    add_hyperlink,
    add_picture_paragraph_after,
    find_placeholders,
    insert_page_break_after,
    insert_paragraph_after,
    insert_table_after,
    load_docx_template,
    normalize_inline_image_sizes,
    remove_review_highlights,
    replace_block_placeholder,
    replace_text_placeholders,
)


BLOCK_PLACEHOLDERS = {"comparables_table", "calculation_table", "report_listing_images", "appendix_full_page_screenshots"}


@dataclass(frozen=True)
class WordReportResult:
    output_path: Path
    replaced_placeholders: set[str]
    missing_placeholders: set[str]
    warnings: list[str]


def generate_word_report(
    *,
    template_path: Path,
    output_path: Path,
    intake_json: Path | None = None,
    candidates_json: Path | None = None,
    task_path: Path | None = None,
    excel_path: Path | None = None,
    include_full_screenshots: bool = False,
) -> WordReportResult:
    template = _resolve_path(template_path)
    output = _resolve_path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = load_docx_template(template)

    intake = _load_intake(_resolve_optional(intake_json))
    task = _load_yaml(_resolve_optional(task_path))
    candidates = _load_candidates(_resolve_optional(candidates_json))
    excel = _resolve_optional(excel_path)

    values = build_report_values(
        intake=intake,
        task=task,
        candidates=candidates,
        excel_path=excel,
        excel_values=read_excel_report_values(excel),
    )
    original_placeholders = find_placeholders(document)
    replaced = replace_text_placeholders(document, values, skip=BLOCK_PLACEHOLDERS)
    warnings: list[str] = []

    if replace_block_placeholder(document, "comparables_table", lambda doc, anchor: _insert_comparables_table(doc, anchor, candidates)):
        replaced.add("comparables_table")
    elif "comparables_table" in original_placeholders:
        warnings.append("comparables_table_placeholder_not_filled")

    if replace_block_placeholder(document, "calculation_table", lambda doc, anchor: _insert_calculation_note(doc, anchor, excel)):
        replaced.add("calculation_table")

    if not _update_existing_comparables_table(document, candidates, values) and candidates:
        warnings.append("existing_comparables_table_not_found")
    if excel and not _update_existing_adjustment_table(document, excel):
        warnings.append("existing_adjustment_table_not_updated")

    if replace_block_placeholder(document, "report_listing_images", lambda doc, anchor: _insert_listing_images(anchor, candidates, full_page=False)):
        replaced.add("report_listing_images")
    elif "report_listing_images" in original_placeholders:
        warnings.append("report_listing_images_placeholder_not_filled")

    if include_full_screenshots:
        if replace_block_placeholder(
            document,
            "appendix_full_page_screenshots",
            lambda doc, anchor: _insert_listing_images(anchor, candidates, full_page=True),
        ):
            replaced.add("appendix_full_page_screenshots")
    else:
        replaced.update(replace_text_placeholders(document, {"appendix_full_page_screenshots": ""}))

    missing = {name for name in original_placeholders if name not in replaced and name not in values}
    if missing:
        warnings.append(f"unresolved_placeholders: {', '.join(sorted(missing))}")

    removed_highlights = remove_review_highlights(document)
    if removed_highlights:
        warnings.append(f"removed_review_highlights: {removed_highlights}")
    resized_images = normalize_inline_image_sizes(document)
    if resized_images:
        warnings.append(f"resized_oversized_images: {resized_images}")

    document.save(str(output))
    # Замінюємо статичні скани витяга/техпаспорта в шаблоні на сторінки ЦЬОГО об'єкта.
    warnings.extend(_swap_object_document_scans(output, intake, task))
    # Замінюємо зразкові скриншоти 5 аналогів на скриншоти відібраних аналогів.
    warnings.extend(_swap_analog_screenshots(output, candidates))
    return WordReportResult(
        output_path=output,
        replaced_placeholders=replaced,
        missing_placeholders=missing,
        warnings=warnings,
    )


# Картинки шаблону valuation_report_real_template, що є сканами документів зразка
# і заміняються сканом ПОТОЧНОГО об'єкта (вит’яг / техпаспорт).
_DOC_SCAN_TARGETS = {"vityag": "word/media/image5.png", "techpassport": "word/media/image1.png"}

# Статичні скриншоти 5 аналогів-зразків у шаблоні (порядок документа = аналог 1→5).
_ANALOG_SLOTS = (
    "word/media/image14.png", "word/media/image13.png", "word/media/image11.png",
    "word/media/image12.png", "word/media/image9.png",
)


def _swap_analog_screenshots(output: Path, candidates: list[Comparable]) -> list[str]:
    """Підставляє скриншоти ВІДІБРАНИХ аналогів замість зразкових у шаблоні
    (аналог i → слот i). Кандидати без скриншота лишають зразок шаблону."""
    notes: list[str] = []
    try:
        mapping: dict[str, Path] = {}
        aspects: dict[str, float] = {}
        for i, cand in enumerate(candidates[: len(_ANALOG_SLOTS)]):
            img = cand.report_image_path or cand.screenshot_path
            path = Path(str(img)) if img else None
            if path and path.exists():
                mapping[_ANALOG_SLOTS[i]] = path
                aspects[_ANALOG_SLOTS[i]] = _image_aspect(path)
        if mapping:
            _replace_docx_media(output, mapping)
            _fit_media_drawings(output, aspects)
            notes.append(f"analog_screens_swapped: {len(mapping)}/{min(len(_ANALOG_SLOTS), len(candidates))}")
    except Exception as exc:  # noqa: BLE001 — підміна скринів не валить звіт
        notes.append(f"analog_screens_failed: {exc}")
    return notes


def _swap_object_document_scans(output: Path, intake: IntakeResult | None, task: dict[str, Any]) -> list[str]:
    """Рендерить сторінку витяга + сторінку техпаспорта об'єкта з PDF і підставляє
    їх замість статичних сканів-зразків у шаблоні (за тим самим розміщенням)."""
    import tempfile

    if intake is None:
        return ["object_scans_skipped_no_intake"]
    src_pdf = _source_pdf(intake, task)
    if not src_pdf or not src_pdf.exists():
        return ["object_scans_no_source_pdf"]

    se = intake.selected_extract
    st = intake.selected_technical_passport
    vityag_page = getattr(se, "page", None) if se else None
    tech_pages_all = list(getattr(st, "pages", []) or []) if st else []
    tech_pages = _pick_techpass_pages(src_pdf, tech_pages_all)

    notes: list[str] = []
    try:
        from realtify.pdf_tools import render_pdf_pages

        with tempfile.TemporaryDirectory(prefix="docscan_") as tmp:
            tmpd = Path(tmp)
            mapping: dict[str, Path] = {}
            aspects: dict[str, float] = {}
            if vityag_page:
                imgs = render_pdf_pages(src_pdf, tmpd / "v", first_page=vityag_page, last_page=vityag_page, dpi=160)
                if imgs:
                    mapping[_DOC_SCAN_TARGETS["vityag"]] = imgs[0]
                    aspects[_DOC_SCAN_TARGETS["vityag"]] = _image_aspect(imgs[0])
            if tech_pages:
                rendered: list[Path] = []
                for idx, page in enumerate(tech_pages):
                    imgs = render_pdf_pages(src_pdf, tmpd / f"t{idx}", first_page=page, last_page=page, dpi=160)
                    if imgs:
                        rendered.append(imgs[0])
                if rendered:
                    stacked, aspect = _stack_images_vertical(rendered, tmpd / "techpass.png")
                    mapping[_DOC_SCAN_TARGETS["techpassport"]] = stacked
                    aspects[_DOC_SCAN_TARGETS["techpassport"]] = aspect
            if not mapping:
                return ["object_scans_no_pages"]
            _replace_docx_media(output, mapping)
            _fit_media_drawings(output, aspects)
            notes.append(f"object_scans_swapped: vityag=p{vityag_page}, techpass=p{tech_pages}")
    except Exception as exc:  # noqa: BLE001 — підміна сканів не повинна валити звіт
        notes.append(f"object_scans_failed: {exc}")
    return notes


def _pick_techpass_pages(pdf: Path, pages: list[int]) -> list[int]:
    """Сторінки техпаспорта для звіту, за текстом (pdftotext):
      • старий формат — сторінка з планом поверху/експлікацією (1 стор.);
      • цифровий формат — інформаційна сторінка (адреса/замовники) + сторінка ТЕП
        «Поверх розташування / Загальна площа» (2 стор.).
    Титул і підписи виключаються."""
    if not pages:
        return []
    if len(pages) == 1:
        return [pages[0]]
    import subprocess

    from realtify.paths import find_poppler_bin

    poppler = find_poppler_bin()
    exe = str(poppler / "pdftotext") if poppler else "pdftotext"
    texts: dict[int, str] = {}
    for page in pages:
        try:
            result = subprocess.run(
                [exe, "-f", str(page), "-l", str(page), str(pdf), "-"],
                capture_output=True, text=True, timeout=30,
            )
            texts[page] = (result.stdout or "").lower()
        except Exception:  # noqa: BLE001
            texts[page] = ""

    plan_kw = ("план поверх", "експлікац")
    tep_kw = (("поверх розташ", 4), ("загальна площа", 2), ("житлова площа", 1), ("кількість житлових кімнат", 1))
    # витяг з Реєстру буд. діяльності — НЕ титул БТІ і не сторінка підписів
    info_kw = (("реєстру будівельної діяльності", 4), ("реєстраційний номер документ", 3), ("єдиної державної електронної", 2))

    def score(text: str, kws: tuple[tuple[str, int], ...]) -> int:
        return sum(w for kw, w in kws if kw in text)

    # Старий формат: сторінка з планом/експлікацією — її достатньо.
    plan = [p for p in pages if any(k in texts[p] for k in plan_kw)]
    if plan:
        return [plan[0]]

    # Цифровий формат: витяг з Реєстру (адреса/замовники) + ТЕП (поверх/площа).
    tep_ranked = sorted(((score(texts[p], tep_kw), p) for p in pages), reverse=True)
    tep_page = tep_ranked[0][1] if tep_ranked and tep_ranked[0][0] > 0 else None
    info_ranked = sorted(((score(texts[p], info_kw), p) for p in pages if p != tep_page), reverse=True)
    info_page = info_ranked[0][1] if info_ranked and info_ranked[0][0] > 0 else None
    result = sorted(p for p in (info_page, tep_page) if p is not None)
    return result or [pages[len(pages) // 2]]


def _image_aspect(path: Path) -> float:
    try:
        from PIL import Image

        with Image.open(path) as im:
            return im.height / im.width if im.width else 1.414
    except Exception:  # noqa: BLE001
        return 1.414


def _stack_images_vertical(paths: list[Path], out: Path) -> tuple[Path, float]:
    """Склеює сторінки вертикально в одне зображення (спільна ширина). Повертає
    (шлях, співвідношення H/W). Для 1 сторінки чи без PIL — без склейки."""
    if len(paths) == 1:
        return paths[0], _image_aspect(paths[0])
    try:
        from PIL import Image
    except Exception:  # noqa: BLE001
        return paths[0], _image_aspect(paths[0])
    gap = 24
    imgs = [Image.open(p).convert("RGB") for p in paths]
    width = max(im.width for im in imgs)
    scaled = [im if im.width == width else im.resize((width, round(im.height * width / im.width))) for im in imgs]
    total_h = sum(im.height for im in scaled) + gap * (len(scaled) - 1)
    canvas = Image.new("RGB", (width, total_h), "white")
    y = 0
    for im in scaled:
        canvas.paste(im, (0, y))
        y += im.height + gap
    canvas.save(out)
    return out, (total_h / width if width else 1.414)


def _fit_media_drawings(docx_path: Path, aspects: dict[str, float]) -> None:
    """Підганяє рамки (drawing extent) підмінених сканів під фактичне співвідношення
    сторін, з обмеженням висоти однією сторінкою (щоб 2-сторінковий ТЕП не
    «розтягувався»). Ширину зберігає, якщо вписується; інакше масштабує пропорційно."""
    import re
    import shutil
    import zipfile

    max_h = 8_000_000  # ~22 см у EMU
    with zipfile.ZipFile(docx_path) as z:
        rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        doc = z.read("word/document.xml").decode("utf-8")

    rid_aspect: dict[str, float] = {}
    for arc, aspect in aspects.items():
        media = arc.split("/")[-1]
        m = re.search(r'Id="([^"]+)"[^>]*Target="[^"]*' + re.escape(media) + r'"', rels) or re.search(
            r'Target="[^"]*' + re.escape(media) + r'"[^>]*Id="([^"]+)"', rels
        )
        if m:
            rid_aspect[m.group(1)] = aspect
    if not rid_aspect:
        return

    def patch(block: str) -> str:
        for rid, aspect in rid_aspect.items():
            if f'r:embed="{rid}"' not in block:
                continue
            me = re.search(r'<wp:extent\b[^>]*\bcx="(\d+)"[^>]*\bcy="\d+"', block)
            if not me:
                return block
            cx = int(me.group(1))
            cy = round(cx * aspect)
            if cy > max_h:
                cy, cx = max_h, round(max_h / aspect)
            block = re.sub(r'(<wp:extent\b[^>]*\bcx=")\d+("[^>]*\bcy=")\d+(")',
                           rf'\g<1>{cx}\g<2>{cy}\g<3>', block, count=1)
            block = re.sub(r'(<a:ext\b[^>]*\bcx=")\d+("[^>]*\bcy=")\d+(")',
                           rf'\g<1>{cx}\g<2>{cy}\g<3>', block, count=1)
            return block
        return block

    doc = re.sub(r"<w:drawing>.*?</w:drawing>", lambda m: patch(m.group(0)), doc, flags=re.S)
    tmp = docx_path.with_suffix(".fit.docx")
    with zipfile.ZipFile(docx_path) as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = doc.encode("utf-8") if item.filename == "word/document.xml" else zin.read(item.filename)
            zout.writestr(item, data)
    shutil.move(str(tmp), str(docx_path))


def _source_pdf(intake: IntakeResult, task: dict[str, Any]) -> Path | None:
    candidates = [getattr(intake, "source_pdf", None)]
    docs = task.get("documents") if isinstance(task, dict) else None
    if isinstance(docs, dict):
        candidates += [docs.get("extract_pdf"), docs.get("technical_passport_pdf")]
    for c in candidates:
        if c:
            p = Path(str(c))
            if p.exists():
                return p
    return None


def _replace_docx_media(docx_path: Path, mapping: dict[str, Path]) -> None:
    """Заміна байтів media-файлів у docx (zip) — той самий arcname, нові байти."""
    import shutil
    import zipfile

    tmp = docx_path.with_suffix(".swap.docx")
    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            repl = mapping.get(item.filename)
            if repl is not None and repl.exists():
                data = repl.read_bytes()
            zout.writestr(item, data)
    shutil.move(str(tmp), str(docx_path))


def _dominant_complex(candidates: list[Comparable]) -> str:
    """Назва ЖК об'єкта = найчастіша назва ЖК серед відібраних аналогів того ж будинку.
    Витяг/техпаспорт назви ЖК не містять, тож беремо її з аналогів того самого дому."""
    from collections import Counter

    names = [
        str(c.complex_name).strip()
        for c in candidates
        if getattr(c, "complex_name", None) and str(c.complex_name).strip().lower() not in ("", "none", "nan")
    ]
    if not names:
        return ""
    return Counter(names).most_common(1)[0][0]


# Видимий плейсхолдер для ручного дозаповнення (даних по об'єкту в движка немає).
_FILL_BLANK = "________________"


def build_report_values(
    *,
    intake: IntakeResult | None,
    task: dict[str, Any],
    candidates: list[Comparable],
    excel_path: Path | None,
    excel_values: dict[str, Any] | None = None,
) -> dict[str, Any]:
    selected_extract = intake.selected_extract if intake else None
    selected_technical = intake.selected_technical_passport if intake else None
    target = task.get("target") if isinstance(task.get("target"), dict) else {}
    documents = task.get("documents") if isinstance(task.get("documents"), dict) else {}

    total_area = _first_not_empty(
        target.get("area_m2"),
        selected_extract.total_area_m2 if selected_extract else None,
        selected_technical.total_area_m2 if selected_technical else None,
    )
    living_area = _first_not_empty(
        target.get("living_area_m2"),
        selected_extract.living_area_m2 if selected_extract else None,
        selected_technical.living_area_m2 if selected_technical else None,
    )
    rooms = _first_not_empty(
        target.get("rooms"),
        selected_technical.rooms_count if selected_technical else None,
    )
    address = _first_not_empty(
        target.get("address"),
        selected_extract.address_full if selected_extract else None,
        selected_technical.address_full if selected_technical else None,
    )
    city = _first_not_empty(
        target.get("city"),
        selected_extract.city if selected_extract else None,
        selected_technical.city if selected_technical else None,
    )
    apartment = _first_not_empty(
        selected_extract.apartment_number if selected_extract else None,
        selected_technical.apartment_number if selected_technical else None,
        _target_apartment_from_address(address),
    )
    address = _normalize_report_address(address)
    excel_payload = excel_values or {}
    valuation_date_obj = resolve_valuation_date(task=task, excel_path=excel_path)
    # Дата складання звіту = дата оцінки (вимога клієнта), якщо явно не задано іншу.
    report_date_obj = _parse_date(_first_not_empty(task.get("report_date"), target.get("report_date"))) or valuation_date_obj
    address_parts = _address_parts(address)
    rooms_text = _rooms_text(rooms)
    total_area_text = _format_decimal_comma(total_area)
    market_value_rounded = excel_payload.get("market_value_uah_rounded")
    market_value_conclusion = _market_value_conclusion(market_value_rounded, excel_payload.get("market_value_uah_words"))

    values: dict[str, Any] = {
        "report_date": report_date_obj.strftime("%d.%m.%Y"),
        "report_date_long": _date_long_uk(report_date_obj),
        "valuation_date": valuation_date_obj.strftime("%d.%m.%Y"),
        "valuation_date_long": _date_long_uk(valuation_date_obj),
        "report_year": str(report_date_obj.year),
        "report_city": city or "",
        "address_full": address or "",
        "address_full_upper": str(address or "").upper(),
        "address_city": city or "",
        "address_line_1_upper": address_parts["line_1_upper"],
        "address_building": address_parts["building"],
        "address_short": address_parts["short"],
        "address_apartment": apartment or "",
        "apartment_number": apartment or "",
        "property_type_text": _property_type_text(target.get("property_type"), rooms),
        "rooms_text": rooms_text,
        "rooms_text_upper": rooms_text.upper() if rooms_text else "",
        "object_type": _first_not_empty(selected_extract.object_type if selected_extract else None, target.get("property_type")),
        "object_description": selected_extract.object_description if selected_extract else "",
        "object_valuation_description": _object_valuation_description(rooms, apartment, total_area, address),
        "object_valuation_description_short": _object_description_short(rooms, apartment, total_area),
        "total_area_m2": total_area_text,
        "total_area_m2_with_unit": f"{total_area_text} кв.м" if total_area_text else "",
        "living_area_m2": _format_decimal_comma(living_area),
        "rooms_count": rooms or "",
        "floor_or_level": _first_not_empty(target.get("floor_or_level"), selected_technical.floor_or_level if selected_technical else None),
        "complex_name": target.get("complex_name") or "",
        "building_class": target.get("building_class") or "",
        "condition": target.get("condition") or "",
        "delivery_date": target.get("delivery_date") or "",
        "extract_index_number": _first_not_empty(
            documents.get("extract_index_number"),
            selected_extract.extract_index_number if selected_extract else None,
        ),
        "extract_date": _extract_date(
            _first_not_empty(documents.get("extract_formed_at"), selected_extract.extract_formed_at if selected_extract else None)
        ),
        "extract_formed_at": _first_not_empty(
            documents.get("extract_formed_at"),
            selected_extract.extract_formed_at if selected_extract else None,
        ),
        "registry_object_number": _first_not_empty(
            documents.get("registry_object_number"),
            selected_extract.registry_object_number if selected_extract else None,
        ),
        "owners_from_extract": _first_not_empty(
            documents.get("owners_from_extract"),
            selected_extract.owners_from_extract if selected_extract else None,
        ),
        "technical_passport_registration_number": _first_not_empty(
            documents.get("technical_passport_registration_number"),
            selected_technical.registration_number if selected_technical else None,
        ),
        "technical_passport_date": _first_not_empty(
            documents.get("technical_passport_date"),
            selected_technical.technical_passport_date if selected_technical else None,
        ),
        "excel_output_path": str(excel_path) if excel_path else "",
        "comparables_count": len(candidates),
        "average_price_usd_m2": _format_decimal_comma(_first_not_empty(excel_payload.get("average_price_usd_m2"), _average_price_m2(candidates))),
        "median_price_usd_m2": _format_decimal_comma(_first_not_empty(excel_payload.get("median_price_usd_m2"), _median_price_m2(candidates))),
        "median_price_usd_m2_report": _format_money_decimal(excel_payload.get("median_price_usd_m2")),
        "nbu_rate": _format_rate(excel_payload.get("nbu_rate")),
        "median_price_uah_m2": _format_money_decimal(excel_payload.get("median_price_uah_m2")),
        "market_value_uah": _format_money_decimal(excel_payload.get("market_value_uah")),
        "market_value_uah_raw": _format_money_decimal(excel_payload.get("market_value_uah")),
        "market_value_uah_rounded": _format_money_decimal(market_value_rounded),
        "market_value_uah_words": excel_payload.get("market_value_uah_words") or "",
        "market_value_uah_conclusion": market_value_conclusion,
    }
    values["extract_reference"] = _extract_reference(values["extract_index_number"], values["extract_date"])
    # Опис ЖК/району об'єкта: назву ЖК беремо з аналогів того ж будинку (авто),
    # а детальні характеристики (клас/поверховість/технологія) та опис району —
    # видимі плейсхолдери для ручного дозаповнення (даних по об'єкту в движка немає).
    values["object_complex_name"] = (
        _first_not_empty(target.get("complex_name"), _dominant_complex(candidates)) or f"{_FILL_BLANK} (заповнити)"
    )
    values["object_building_details"] = (
        f"Клас будинку — {_FILL_BLANK}; кількість будинків — ________; поверховість — ________; "
        f"технологія будівництва — {_FILL_BLANK}; опалення — {_FILL_BLANK} (заповнити вручну)."
    )
    values["location_description"] = (
        f"{_FILL_BLANK}{_FILL_BLANK} (характеристика місцезнаходження та району розташування "
        "об'єкта оцінки — заповнити вручну)."
    )
    for idx, candidate in enumerate(candidates[:5], start=1):
        prefix = f"comparable_{idx}_"
        values.update(
            {
                f"{prefix}address": candidate.address or "",
                f"{prefix}area_m2": _format_number(candidate.area_m2),
                f"{prefix}floor_or_level": candidate.floor_or_level or "",
                f"{prefix}price_usd": _format_money(candidate.price_usd),
                f"{prefix}price_usd_m2": _format_number(candidate.price_per_m2_usd),
                f"{prefix}location_quality": candidate.location_quality or "",
                f"{prefix}building_class": candidate.building_class or "",
                f"{prefix}condition": candidate.condition or "",
                f"{prefix}delivery_date": candidate.delivery_date or "",
                f"{prefix}source_url": str(candidate.source_url),
            }
        )
    return values


def _insert_comparables_table(document, anchor, candidates: list[Comparable]) -> None:
    if not candidates:
        insert_paragraph_after(anchor, "Аналоги не зібрано.")
        return
    headers = ["#", "Адреса", "Площа", "Поверх", "Ціна, USD", "USD/кв.м", "Джерело"]
    table = insert_table_after(document, anchor, rows=len(candidates[:5]) + 1, cols=len(headers))
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row_idx, candidate in enumerate(candidates[:5], start=1):
        values = [
            str(row_idx),
            candidate.address or "",
            _format_number(candidate.area_m2),
            candidate.floor_or_level or "",
            _format_money(candidate.price_usd),
            _format_number(candidate.price_per_m2_usd),
            "",
        ]
        for col_idx, value in enumerate(values):
            table.cell(row_idx, col_idx).text = value
        paragraph = table.cell(row_idx, 6).paragraphs[0]
        add_hyperlink(paragraph, str(candidate.source_url), "посилання")


def _insert_calculation_note(document, anchor, excel_path: Path | None) -> None:
    paragraph = insert_paragraph_after(anchor)
    if excel_path:
        paragraph.add_run("Розрахункова таблиця сформована у файлі: ")
        paragraph.add_run(str(excel_path))
    else:
        paragraph.add_run("Розрахункова таблиця не передана в генератор Word-звіту.")


def _update_existing_comparables_table(document, candidates: list[Comparable], values: dict[str, Any]) -> bool:
    updated = False
    for table in document.tables:
        if not _is_existing_comparables_table(table):
            continue
        _fill_existing_comparables_table(table, candidates[:5], values)
        updated = True
    return updated


def _is_existing_comparables_table(table) -> bool:
    if len(table.rows) < 11 or len(table.columns) < 7:
        return False
    header_text = " ".join(cell.text for cell in table.rows[0].cells[:7]).casefold()
    first_col_text = " ".join(row.cells[0].text for row in table.rows[:11]).casefold()
    return (
        "об'єкт порівняння №1" in header_text
        and "об'єкт оцінки" in header_text
        and "адреса об'єкта порівняння" in first_col_text
        and "ціна пропозиції" in first_col_text
        and "джерело інформації" in first_col_text
    )


def _fill_existing_comparables_table(table, candidates: list[Comparable], values: dict[str, Any]) -> None:
    defaults = _target_defaults(values)
    for idx in range(5):
        col = idx + 1
        candidate = candidates[idx] if idx < len(candidates) else None
        _write_table_column(table, col, _candidate_table_values(candidate, defaults))

    _write_table_column(
        table,
        6,
        {
            1: _target_table_address(values),
            2: str(values.get("total_area_m2") or ""),
            3: str(values.get("floor_or_level") or ""),
            4: "",
            5: "",
            6: defaults["location_quality"],
            7: defaults["building_class"],
            8: defaults["condition"],
            9: defaults["delivery_date"],
            10: "Х",
        },
    )


def _target_defaults(values: dict[str, Any]) -> dict[str, str]:
    return {
        "location_quality": str(values.get("location_quality") or "добре"),
        "building_class": str(values.get("building_class") or "Комфорт"),
        "condition": str(values.get("condition") or "від забудовника"),
        "delivery_date": str(values.get("delivery_date") or "введений"),
    }


def _candidate_table_values(candidate: Comparable | None, defaults: dict[str, str]) -> dict[int, str]:
    if not candidate:
        return {row: "" for row in range(1, 11)}
    return {
        1: _candidate_table_address(candidate),
        2: _format_decimal_comma(candidate.area_m2),
        3: candidate.floor_or_level or "",
        4: _format_money(candidate.price_usd),
        5: _format_money(candidate.price_per_m2_usd),
        6: candidate.location_quality or defaults["location_quality"],
        7: candidate.building_class or defaults["building_class"],
        8: candidate.condition or defaults["condition"],
        9: candidate.delivery_date or defaults["delivery_date"],
        10: str(candidate.source_url),
    }


def _write_table_column(table, col: int, row_values: dict[int, str]) -> None:
    for row_offset, value in row_values.items():
        cell = table.cell(row_offset, col)
        if row_offset == 10 and value.startswith(("http://", "https://")):
            _set_cell_hyperlink(cell, value)
        else:
            _set_cell_text(cell, value)


def _set_cell_text(cell, value: str) -> None:
    cell.text = value


def _set_cell_hyperlink(cell, url: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    add_hyperlink(paragraph, url, url)


def _candidate_table_address(candidate: Comparable) -> str:
    parts = [candidate.address or ""]
    if candidate.complex_name and candidate.complex_name not in parts[0]:
        parts.append(candidate.complex_name)
    return ", ".join(part for part in parts if part)


def _target_table_address(values: dict[str, Any]) -> str:
    address = str(values.get("address_short") or values.get("address_full") or "")
    complex_name = str(values.get("complex_name") or "")
    if complex_name and complex_name not in address:
        return f"{address}, {complex_name}" if address else complex_name
    return address


def _update_existing_adjustment_table(document, excel_path_value: Path) -> bool:
    excel_rows = _read_adjustment_rows_from_excel(excel_path_value)
    if not excel_rows:
        return False
    updated = False
    for table in document.tables:
        if not _is_existing_adjustment_table(table):
            continue
        _fill_existing_adjustment_table(table, excel_rows)
        updated = True
    return updated


def _is_existing_adjustment_table(table) -> bool:
    if len(table.rows) < 27 or len(table.columns) < 8:
        return False
    first_col_text = " ".join(row.cells[0].text for row in table.rows[:27]).casefold()
    header_text = " ".join(cell.text for cell in table.rows[0].cells[:8]).casefold()
    return (
        "об’єкт оцінки" in header_text
        and "об’єкт порівняння №1" in header_text
        and "ціна 1 м2 загальної площі" in first_col_text
        and "вартість з урахуванням поправок" in first_col_text
        and "медіана варіаційного ряду" in first_col_text
    )


def _fill_existing_adjustment_table(table, excel_rows: dict[int, list[str]]) -> None:
    word_to_excel_rows = [
        15,
        16,
        17,
        18,
        19,
        20,
        23,
        24,
        25,
        26,
        27,
        28,
        29,
        30,
        31,
        32,
        33,
        34,
        35,
        36,
        37,
        38,
        39,
        40,
        41,
        42,
        43,
    ]
    table_rows = table.rows
    for word_row_index, excel_row_index in enumerate(word_to_excel_rows):
        row_values = excel_rows.get(excel_row_index)
        if not row_values:
            continue
        if word_row_index >= len(table_rows):
            break
        # КРИТИЧНО: тримаємо посилання на ВСІ клітинки рядка (list(...)) на весь
        # внутрішній цикл. Інакше прокси _tc попередньої клітинки звільняє GC, а
        # id() наступної переиспользует звільнену адресу → хибне спрацювання
        # seen_cells → клітинку пропускає й лишається шаблонне значення (баг
        # недетермінований, проявляється під навантаженням повного пайплайна).
        cells = list(table_rows[word_row_index].cells)
        seen_cells: set[int] = set()
        for col_index, value in enumerate(row_values[:8]):
            if col_index >= len(cells):
                break
            cell = cells[col_index]
            cell_id = id(cell._tc)
            if cell_id in seen_cells:
                continue
            seen_cells.add(cell_id)
            _set_cell_text(cell, value)


def _read_adjustment_rows_from_excel(path: Path) -> dict[int, list[str]]:
    sidecar_rows = sidecar_adjustment_rows(path)
    if sidecar_rows:
        return sidecar_rows
    try:
        import pythoncom
        import win32com.client
    except Exception as exc:
        return {}

    rows: dict[int, list[str]] = {}
    pythoncom.CoInitialize()
    excel = win32com.client.DispatchEx("Excel.Application")
    excel.Visible = False
    excel.DisplayAlerts = False
    try:
        wb = excel.Workbooks.Open(excel_path(path), 0, True)
        try:
            ws = _find_calculation_sheet(wb)
            for row_index in range(15, 44):
                rows[row_index] = [str(ws.Cells(row_index, col_index).Text) for col_index in range(1, 9)]
        finally:
            wb.Close(False)
    finally:
        excel.Quit()
        pythoncom.CoUninitialize()
    return rows


def _find_calculation_sheet(workbook):
    for index in range(1, workbook.Worksheets.Count + 1):
        sheet = workbook.Worksheets(index)
        if "Розрахунок" in str(sheet.Name):
            return sheet
    return workbook.Worksheets(1)


def _insert_listing_images(anchor, candidates: list[Comparable], *, full_page: bool) -> None:
    current = anchor
    for index, candidate in enumerate(candidates[:5], start=1):
        image = candidate.screenshot_path if full_page else candidate.report_image_path
        if not image:
            image = candidate.screenshot_path
        image_path = Path(image) if image else None
        caption = f"Аналог {index}. {candidate.address or candidate.title or 'Оголошення'}"
        if image_path and image_path.exists():
            current = add_picture_paragraph_after(current, image_path, caption=caption)
        else:
            current = insert_paragraph_after(current, caption)
            current.add_run(" - скріншот не знайдено.")
        link_para = insert_paragraph_after(current)
        link_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for run in link_para.runs:
            run.font.size = Pt(9)
        add_hyperlink(link_para, str(candidate.source_url), str(candidate.source_url))
        if full_page and index < min(5, len(candidates)):
            current = insert_page_break_after(link_para)
        else:
            current = link_para


def _load_intake(path: Path | None) -> IntakeResult | None:
    if not path:
        return None
    with path.open("r", encoding="utf-8") as f:
        return IntakeResult.model_validate(json.load(f))


def _load_candidates(path: Path | None) -> list[Comparable]:
    if not path:
        return []
    with path.open("r", encoding="utf-8") as f:
        payload = json.load(f)
    raw_candidates = payload.get("candidates", [])
    try:
        return [Comparable.model_validate(item) for item in raw_candidates]
    except ValidationError as exc:
        raise WordTemplateError(f"Invalid candidates JSON: {exc}") from exc


def _load_yaml(path: Path | None) -> dict[str, Any]:
    if not path:
        return {}
    with path.open("r", encoding="utf-8") as f:
        payload = yaml.safe_load(f) or {}
    if not isinstance(payload, dict):
        raise WordTemplateError(f"{path} must contain a YAML object")
    return payload


def _resolve_optional(path: Path | None) -> Path | None:
    if not path:
        return None
    return _resolve_path(path)


def _resolve_path(path: Path) -> Path:
    return path if path.is_absolute() else PROJECT_ROOT / path


def _first_not_empty(*values: Any) -> Any:
    for value in values:
        if value not in (None, ""):
            return value
    return ""


def _format_number(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    if number.is_integer():
        return str(int(number))
    return f"{number:.2f}".rstrip("0").rstrip(".")


def _format_money(value: Any) -> str:
    text = _format_number(value)
    if not text:
        return ""
    try:
        return f"{float(text):,.0f}".replace(",", " ")
    except ValueError:
        return text


def _format_decimal_comma(value: Any, *, decimals: int = 2, trim: bool = True) -> str:
    if value in (None, ""):
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    text = f"{number:.{decimals}f}".replace(".", ",")
    if trim and "," in text:
        text = text.rstrip("0").rstrip(",")
    return text


def _format_money_decimal(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    whole, frac = f"{number:,.2f}".split(".")
    return f"{whole.replace(',', ' ')},{frac}"


def _format_rate(value: Any) -> str:
    if value in (None, ""):
        return ""
    try:
        return f"{float(value):.4f}".replace(".", ",")
    except (TypeError, ValueError):
        return str(value)


def _average_price_m2(candidates: list[Comparable]) -> float | None:
    values = [candidate.price_per_m2_usd for candidate in candidates if candidate.price_per_m2_usd is not None]
    if not values:
        return None
    return sum(values) / len(values)


def _median_price_m2(candidates: list[Comparable]) -> float | None:
    values = sorted(candidate.price_per_m2_usd for candidate in candidates if candidate.price_per_m2_usd is not None)
    if not values:
        return None
    mid = len(values) // 2
    if len(values) % 2:
        return values[mid]
    return (values[mid - 1] + values[mid]) / 2


def _extract_date(value: Any) -> str:
    text = str(value or "")
    return text.split()[0] if text else ""


def _extract_reference(index_number: Any, extract_date: Any) -> str:
    index = str(index_number or "").strip()
    date_text = str(extract_date or "").strip()
    if index and date_text:
        return f"\u2116 {index} \u0432\u0456\u0434 {date_text} \u0440."
    if index:
        return f"\u2116 {index}"
    return ""


def _parse_date(value: Any) -> date | None:
    if not value:
        return None
    text = str(value).strip().split()[0].replace(",", ".")
    for fmt in ("%d.%m.%Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(text, fmt).date()
        except ValueError:
            pass
    return None


def _date_long_uk(value: date) -> str:
    months = [
        "",
        "\u0441\u0456\u0447\u043d\u044f",
        "\u043b\u044e\u0442\u043e\u0433\u043e",
        "\u0431\u0435\u0440\u0435\u0437\u043d\u044f",
        "\u043a\u0432\u0456\u0442\u043d\u044f",
        "\u0442\u0440\u0430\u0432\u043d\u044f",
        "\u0447\u0435\u0440\u0432\u043d\u044f",
        "\u043b\u0438\u043f\u043d\u044f",
        "\u0441\u0435\u0440\u043f\u043d\u044f",
        "\u0432\u0435\u0440\u0435\u0441\u043d\u044f",
        "\u0436\u043e\u0432\u0442\u043d\u044f",
        "\u043b\u0438\u0441\u0442\u043e\u043f\u0430\u0434\u0430",
        "\u0433\u0440\u0443\u0434\u043d\u044f",
    ]
    return f"{value.day} {months[value.month]} {value.year} \u0440\u043e\u043a\u0443"


def _rooms_text(value: Any) -> str:
    try:
        rooms = int(value)
    except (TypeError, ValueError):
        return ""
    mapping = {1: "однокімнатної", 2: "двокімнатної", 3: "трикімнатної", 4: "чотирикімнатної"}
    return mapping.get(rooms, f"{rooms}-кімнатної")


def _property_type_text(property_type: Any, rooms: Any) -> str:
    if property_type == "apartment":
        rooms_text = _rooms_text(rooms)
        return f"{rooms_text} квартири" if rooms_text else "квартири"
    if property_type == "parking":
        return "машиномісця у підземному паркінгу"
    return str(property_type or "об'єкта нерухомості")


def _target_apartment_from_address(address: Any) -> str:
    if not address:
        return ""
    import re

    match = re.search(r"(?:квартира|кв\.?/оф\.?|кв\.)\s*([0-9]{1,5})", str(address), re.IGNORECASE)
    return match.group(1) if match else ""


def _address_parts(address: Any) -> dict[str, str]:
    import re

    text = str(address or "").strip()
    clean = " ".join(text.replace("\u00a0", " ").split())
    clean = re.sub(r"\bм\.\s*", "м. ", clean, flags=re.IGNORECASE)
    building = ""
    match = re.search(r"(?:будинок|буд\.?)\s*([0-9A-Za-zА-Яа-яІіЇїЄєҐґ\\-]+)", clean, re.IGNORECASE)
    if match:
        building = match.group(1)
    without_apartment = re.sub(r",?\s*(?:квартира|кв\.?/оф\.?|кв\.)\s*\d+.*$", "", clean, flags=re.IGNORECASE)
    before_building = re.split(r",?\s*будинок\b", without_apartment, flags=re.IGNORECASE)[0].strip(" ,")
    return {
        "line_1_upper": before_building.upper() + ("," if before_building and not before_building.endswith(",") else ""),
        "building": building,
        "short": without_apartment,
    }


def _normalize_report_address(address: Any) -> str:
    import re

    clean = " ".join(str(address or "").replace("\u00a0", " ").split())
    clean = re.sub(r"\bм\.\s*", "м. ", clean, flags=re.IGNORECASE)
    clean = re.sub(r"\s*,\s*", ", ", clean)
    return clean.strip(" ,")


def _object_description_short(rooms: Any, apartment: Any, total_area: Any) -> str:
    rooms_text = _rooms_text(rooms).capitalize()
    area = _format_decimal_comma(total_area)
    number = str(apartment or "")
    if rooms_text and number and area:
        return f"{rooms_text} квартири № {number} загальною площею {area} кв.м"
    return ""


def _object_valuation_description(rooms: Any, apartment: Any, total_area: Any, address: Any) -> str:
    rooms_text = _rooms_text(rooms)
    area = _format_decimal_comma(total_area)
    number = str(apartment or "")
    address_text = str(address or "")
    if rooms_text and number and area and address_text:
        return (
            f"{rooms_text} квартири № {number}, "
            f"загальною площею {area} кв.м., "
            f"яка розташована в житловому будинку за адресою: {address_text}"
        )
    return ""


def _market_value_conclusion(value: Any, words: Any) -> str:
    if value in (None, ""):
        return ""
    amount = _format_money_decimal(value)
    words_text = str(words or "").strip()
    if words_text:
        return f"{amount} ({words_text}) грн., без ПДВ."
    return f"{amount} грн., без ПДВ."


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Generate a .docx valuation report from a .docx template.")
    parser.add_argument("--template", type=Path, required=True, help="Prepared .docx report template.")
    parser.add_argument("--out", type=Path, default=None, help="Output .docx path.")
    parser.add_argument("--intake", type=Path, default=None, help="intake.json from PDF extraction.")
    parser.add_argument("--candidates", type=Path, default=None, help="candidates.json from listing collection.")
    parser.add_argument("--task", type=Path, default=None, help="task YAML generated from intake.")
    parser.add_argument("--excel", type=Path, default=None, help="Filled Excel workbook path.")
    parser.add_argument("--include-full-screenshots", action="store_true")
    args = parser.parse_args(argv)

    console = Console()
    output = args.out or ensure_output_dir(datetime.now().strftime("%Y%m%d_%H%M%S_word_report")) / "valuation_report.docx"
    try:
        result = generate_word_report(
            template_path=args.template,
            output_path=output,
            intake_json=args.intake,
            candidates_json=args.candidates,
            task_path=args.task,
            excel_path=args.excel,
            include_full_screenshots=args.include_full_screenshots,
        )
    except Exception as exc:
        console.print(f"[red]Word report generation failed:[/red] {exc}")
        return 1

    console.print("[green]Word report generated[/green]")
    if result.warnings:
        for warning in result.warnings:
            console.print(f"[yellow]Warning:[/yellow] {warning}")
    console.print(f"Output: {result.output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
