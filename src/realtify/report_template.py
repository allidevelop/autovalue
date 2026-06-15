from __future__ import annotations

import argparse
import json
import shutil
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from rich.console import Console

from realtify.paths import PROJECT_ROOT, ensure_output_dir
from realtify.word_tools import (
    PLACEHOLDER_RE,
    WordTemplateError,
    apply_basic_business_styles,
    convert_doc_to_docx,
    iter_paragraph_locations,
)


@dataclass(frozen=True)
class PreparedTemplate:
    docx_path: Path
    inventory_json: Path
    inventory_md: Path


def prepare_report_template(source_path: Path, output_dir: Path) -> PreparedTemplate:
    source = _resolve_path(source_path)
    output_dir = _resolve_path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_path = output_dir / f"{source.stem}.template.docx"
    if source.suffix.lower() == ".docx":
        shutil.copy2(source, docx_path)
    elif source.suffix.lower() == ".doc":
        convert_doc_to_docx(source, docx_path)
    else:
        raise WordTemplateError(f"Unsupported report template extension: {source.suffix}")

    inventory = inspect_report_template(docx_path)
    inventory_json = output_dir / "template_inventory.json"
    inventory_md = output_dir / "template_inventory.md"
    inventory_json.write_text(json.dumps(inventory, ensure_ascii=False, indent=2), encoding="utf-8")
    inventory_md.write_text(_inventory_markdown(docx_path, inventory), encoding="utf-8")
    return PreparedTemplate(docx_path=docx_path, inventory_json=inventory_json, inventory_md=inventory_md)


def inspect_report_template(docx_path: Path) -> dict:
    document = Document(str(docx_path))
    items = []
    for location in iter_paragraph_locations(document):
        paragraph = location.paragraph
        text = paragraph.text.strip()
        placeholders = sorted(set(PLACEHOLDER_RE.findall(text)))
        highlighted = [
            run.text.strip()
            for run in paragraph.runs
            if run.font.highlight_color is not None and run.text.strip()
        ]
        if placeholders or highlighted:
            items.append(
                {
                    "area": location.area,
                    "paragraph_index": location.index,
                    "text": text,
                    "placeholders": placeholders,
                    "highlighted_runs": highlighted,
                }
            )
    return {
        "created_at": datetime.now().isoformat(timespec="seconds"),
        "docx_path": str(docx_path),
        "items": items,
        "placeholders": sorted({placeholder for item in items for placeholder in item["placeholders"]}),
        "highlighted_count": sum(len(item["highlighted_runs"]) for item in items),
    }


def create_default_report_template(output_path: Path) -> Path:
    output = _resolve_path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    apply_basic_business_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("ЗВІТ ПРО ОЦІНКУ МАЙНА")
    run.bold = True
    run.font.size = Pt(16)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle.add_run("{{property_type_text}} за адресою: {{address_full}}")

    document.add_heading("Вихідні дані", level=1)
    table = document.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    rows = [
        ("Дата оцінки", "{{valuation_date}}"),
        ("Дата звіту", "{{report_date}}"),
        ("Індексний номер витягу", "{{extract_index_number}}"),
        ("Дата формування витягу", "{{extract_formed_at}}"),
        ("Загальна площа, кв.м", "{{total_area_m2}}"),
        ("Житлова площа, кв.м", "{{living_area_m2}}"),
        ("Кількість кімнат", "{{rooms_count}}"),
        ("Поверх", "{{floor_or_level}}"),
    ]
    for row, (label, value) in zip(table.rows, rows):
        row.cells[0].text = label
        row.cells[1].text = value

    document.add_heading("Об'єкти порівняння", level=1)
    document.add_paragraph("{{comparables_table}}")

    document.add_heading("Розрахункові матеріали", level=1)
    document.add_paragraph("Excel-розрахунок: {{excel_output_path}}")
    document.add_paragraph("{{calculation_table}}")

    document.add_heading("Скріншоти оголошень", level=1)
    document.add_paragraph("{{report_listing_images}}")

    document.save(str(output))
    return output


def _inventory_markdown(docx_path: Path, inventory: dict) -> str:
    lines = [
        "# Report Template Inventory",
        "",
        f"Template: {docx_path}",
        f"Placeholders: {len(inventory['placeholders'])}",
        f"Highlighted runs: {inventory['highlighted_count']}",
        "",
    ]
    if inventory["placeholders"]:
        lines.extend(["## Placeholders", ""])
        lines.extend(f"- `{{{{{placeholder}}}}}`" for placeholder in inventory["placeholders"])
        lines.append("")
    if inventory["items"]:
        lines.extend(["## Locations", ""])
        for item in inventory["items"]:
            lines.extend(
                [
                    f"### {item['area']} paragraph {item['paragraph_index']}",
                    "",
                    item["text"] or "_empty paragraph_",
                    "",
                ]
            )
            if item["highlighted_runs"]:
                lines.append("Highlighted:")
                lines.extend(f"- {value}" for value in item["highlighted_runs"])
                lines.append("")
            if item["placeholders"]:
                lines.append("Placeholders:")
                lines.extend(f"- `{{{{{value}}}}}`" for value in item["placeholders"])
                lines.append("")
    return "\n".join(lines).strip() + "\n"


def _resolve_path(path: Path) -> Path:
    return path if path.is_absolute() else PROJECT_ROOT / path


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Prepare or create a .docx report template.")
    parser.add_argument("--source", type=Path, default=None, help="Existing .doc/.docx report sample.")
    parser.add_argument("--out", type=Path, default=None, help="Output directory for prepared template.")
    parser.add_argument("--create-default", type=Path, default=None, help="Create a basic placeholder .docx template.")
    args = parser.parse_args(argv)

    console = Console()
    try:
        if args.create_default:
            path = create_default_report_template(args.create_default)
            console.print(f"[green]Default report template created[/green]: {path}")
            return 0
        if not args.source:
            raise WordTemplateError("Use --source to prepare an existing template or --create-default to create one.")
        out_dir = args.out or ensure_output_dir(datetime.now().strftime("%Y%m%d_%H%M%S_report_template"))
        result = prepare_report_template(args.source, out_dir)
    except Exception as exc:
        console.print(f"[red]Report template preparation failed:[/red] {exc}")
        return 1

    console.print("[green]Report template prepared[/green]")
    console.print(f"DOCX: {result.docx_path}")
    console.print(f"Inventory JSON: {result.inventory_json}")
    console.print(f"Inventory MD: {result.inventory_md}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
